"""
Aurelinx Intelligence Chat endpoints
Persistent sessions, messages, uploads, and tool-enabled agent responses.
"""

from datetime import datetime
from datetime import timedelta
import hashlib
import asyncio
import json
import os
import re
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from uuid import UUID, uuid4

import httpx
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import func
from sqlmodel import Session, select

from app.core.logging_config import get_logger
from app.core.data_policy import filter_real_records
from app.core.provider_utils import (
    build_local_provider_base_candidates,
    normalize_local_provider_base,
)
from app.core.security import TokenData, get_current_user
from app.models.database import (
    AuditLogTable,
    CandidateTable,
    ChatAttachmentTable,
    ChatMessageTable,
    ChatSessionTable,
    CanonicalCandidateTable,
    CanonicalEmployeeTable,
    EmployeeTable,
    SkillTable,
    IntegrationConnectionTable,
    CompliancePolicyTable,
    InterventionTable,
    ExperienceTable,
    GoldMetricSnapshotTable,
    MLModelCardTable,
    MLDriftSnapshotTable,
    ReleaseGateTable,
    DRRunbookTable,
    ProcurementArtifactTable,
    RawEventTable,
    QuarantineEventTable,
    ConnectorSyncJobTable,
    WorkflowRunTable,
    WorkflowEventTable,
    WorkflowApprovalTable,
    engine,
    get_session,
)
from app.schemas.schemas import (
    ChatAttachmentOut,
    ChatBulkDeleteRequest,
    ChatMessageCreate,
    ChatMessageOut,
    ChatResponse,
    ChatSessionCreate,
    ChatSessionOut,
    ChatSessionRename,
)
from app.workflows.events import (
    ToolEventScope,
    create_workflow_run,
    emit_workflow_event,
    safe_summary,
    update_workflow_run,
    workflow_event_dict,
)

router = APIRouter(prefix="/chat", tags=["chat"])
logger = get_logger(__name__)

UPLOAD_ROOT = Path(os.getenv("CHAT_UPLOAD_ROOT", "/app/data/chat")).resolve()
UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)


def _json_dumps(payload: Dict) -> str:
    """Serialize SSE payloads safely, including UUID/datetime values."""
    return json.dumps(payload, default=str)


try:
    from pypdf import PdfReader  # type: ignore
except Exception:
    PdfReader = None

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None

try:
    import pytesseract  # type: ignore
    from PIL import Image  # type: ignore
except Exception:
    pytesseract = None
    Image = None


def _to_session_out(s: ChatSessionTable) -> ChatSessionOut:
    return ChatSessionOut(
        id=s.id,
        user_id=s.user_id,
        title=s.title,
        created_at=s.created_at,
        updated_at=s.updated_at,
    )


def _get_session_normalized(db: Session, session_id: UUID) -> "ChatSessionTable | None":
    """Lookup a ChatSession by ID, handling both hyphenated and hex (no-hyphen) UUID
    storage formats from older string-based records.
    """
    sid_str = str(session_id)  # '74013167-aac8-...' (hyphenated)
    sid_hex = sid_str.replace("-", "")  # '74013167aac8...'  (no hyphens)

    # Try hyphenated string first (new sessions stored this way)
    row = db.exec(
        select(ChatSessionTable).where(ChatSessionTable.id == sid_str)
    ).first()
    if row:
        return row

    # Try hex (no hyphens) for older sessions
    return db.exec(
        select(ChatSessionTable).where(ChatSessionTable.id == sid_hex)
    ).first()


def _to_message_out(m: ChatMessageTable) -> ChatMessageOut:
    workflow_run_id = None
    workflow_events = []
    try:
        trace = json.loads(m.tool_trace or "{}")
        workflow_run_id = trace.get("workflow_run_id")
        workflow_events = trace.get("workflow_events") or []
    except Exception:
        pass
    return ChatMessageOut(
        id=m.id,
        session_id=m.session_id,
        role=m.role,
        content=m.content,
        tool_trace=m.tool_trace,
        created_at=m.created_at,
        workflow_run_id=workflow_run_id,
        workflow_events=workflow_events,
    )


def _to_attachment_out(a: ChatAttachmentTable) -> ChatAttachmentOut:
    return ChatAttachmentOut(
        id=a.id,
        session_id=a.session_id,
        message_id=a.message_id,
        original_name=a.original_name,
        content_type=a.content_type,
        file_path=a.file_path,
        file_size=a.file_size,
        parsing_status=a.parsing_status,
        parsing_error=a.parsing_error,
        created_at=a.created_at,
    )


def _assert_session_owner(session_row: ChatSessionTable, current_user: TokenData):
    if current_user.is_admin:
        return
    if str(session_row.user_id) == "00000000-0000-0000-0000-000000000000":
        return
    if session_row.user_id != current_user.user_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN, detail="Forbidden session access"
        )


def _write_audit(
    db: Session,
    current_user: TokenData,
    action: str,
    resource_type: str,
    resource_id: str,
    details: Dict,
):
    audit = AuditLogTable(
        user_id=None,  # token user id is string in this project, keep optional DB FK null-safe
        action=action,
        resource_type=resource_type,
        resource_id=None,
        details=json.dumps(details),
        ip_address="127.0.0.1",
    )
    db.add(audit)


def _tokenize(text: str) -> List[str]:
    return re.findall(r"[a-z0-9\+\#\.]{2,}", text.lower())


def _search_employees(
    db: Session, query_text: str, limit: int = 5
) -> List[EmployeeTable]:
    employees = db.exec(select(EmployeeTable)).all()
    query_lower = query_text.lower()
    scored = []
    
    # Identify query intent flags
    wants_at_risk = any(k in query_lower for k in ["risk", "flight", "hazard", "leaving", "attrition"])
    wants_high_morale = any(k in query_lower for k in ["high morale", "happy", "satisfied", "good sentiment"])
    wants_low_morale = any(k in query_lower for k in ["low morale", "unhappy", "dissatisfied", "bad sentiment"])
    
    for emp in employees:
        score = 0.0
        
        # 1. Direct substring checks (name, role, department)
        if emp.full_name and emp.full_name.lower() in query_lower:
            score += 10.0
        for token in query_lower.split():
            if len(token) > 2:
                if emp.full_name and token in emp.full_name.lower():
                    score += 5.0
                if emp.role and token in emp.role.lower():
                    score += 3.0
                if emp.department and token in emp.department.lower():
                    score += 4.0
                    
        # 2. Intent matching
        if wants_at_risk:
            if emp.is_at_risk:
                score += 15.0
            score += (1.0 - (emp.sentiment_score or 0.5)) * 10.0
            if emp.retention_prob is not None:
                score += (1.0 - emp.retention_prob) * 10.0
        elif wants_low_morale:
            score += (1.0 - (emp.sentiment_score or 0.5)) * 15.0
        elif wants_high_morale:
            score += (emp.sentiment_score or 0.5) * 15.0
            
        # 3. Skills overlap
        skills = db.exec(select(SkillTable).where(SkillTable.employee_id == emp.id)).all()
        for s in skills:
            if s.name and s.name.lower() in query_lower:
                score += 5.0
                
        scored.append((score, emp))
        
    scored.sort(key=lambda x: x[0], reverse=True)
    return [row[1] for row in scored[:limit]]


def _search_candidates(
    db: Session, query_text: str, limit: int = 5
) -> List[CandidateTable]:
    candidates = db.exec(select(CandidateTable)).all()
    query_lower = query_text.lower()
    scored = []
    
    # Identify query intent flags
    wants_high_match = any(k in query_lower for k in ["best", "top", "highly matched", "match", "scout"])
    wants_high_morale = any(k in query_lower for k in ["high morale", "happy", "good sentiment"])
    wants_low_morale = any(k in query_lower for k in ["low morale", "unhappy", "bad sentiment"])
    
    for c in candidates:
        score = 0.0
        
        # 1. Direct substring checks
        if c.full_name and c.full_name.lower() in query_lower:
            score += 10.0
        for token in query_lower.split():
            if len(token) > 2:
                if c.full_name and token in c.full_name.lower():
                    score += 5.0
                if c.role and token in c.role.lower():
                    score += 3.0
                if c.department and token in c.department.lower():
                    score += 4.0
                    
        # 2. Intent matching
        if wants_high_match:
            score += (c.match_score or 0.0) * 15.0
        if wants_low_morale:
            score += (1.0 - (c.sentiment_score or 0.5)) * 10.0
        elif wants_high_morale:
            score += (c.sentiment_score or 0.5) * 10.0
            
        # 3. Skills overlap
        skills = db.exec(select(SkillTable).where(SkillTable.candidate_id == c.id)).all()
        for s in skills:
            if s.name and s.name.lower() in query_lower:
                score += 5.0
                
        scored.append((score, c))
        
    scored.sort(key=lambda x: x[0], reverse=True)
    return [row[1] for row in scored[:limit]]


def _compute_dashboard_snapshot(db: Session) -> Dict:
    employees = db.exec(select(EmployeeTable)).all()
    total = len(employees)
    at_risk = len([e for e in employees if e.is_at_risk])
    avg_morale = (
        round(sum([e.sentiment_score for e in employees]) / total, 3) if total else 0.0
    )
    dept_counts = {}
    for e in employees:
        dept_counts[e.department] = dept_counts.get(e.department, 0) + 1
    return {
        "total_workforce": total,
        "at_risk": at_risk,
        "avg_morale": avg_morale,
        "departments": dept_counts,
    }


def _department_risk_summary(employees: List[EmployeeTable]) -> List[Dict[str, Any]]:
    dept_map: Dict[str, int] = {}
    dept_risk: Dict[str, int] = {}
    for e in employees:
        dept = e.department or "Unknown"
        dept_map[dept] = dept_map.get(dept, 0) + 1
        if e.is_at_risk:
            dept_risk[dept] = dept_risk.get(dept, 0) + 1

    summary: List[Dict[str, Any]] = []
    for dept, total in sorted(dept_map.items(), key=lambda item: (-item[1], item[0])):
        at_risk = dept_risk.get(dept, 0)
        summary.append(
            {
                "department": dept,
                "total": total,
                "at_risk": at_risk,
                "at_risk_rate": round((at_risk / total), 3) if total else 0.0,
            }
        )
    return summary


def _top_risk_score_employee(employee: EmployeeTable) -> float:
    sentiment = float(employee.sentiment_score or 0.0)
    retention = float(
        employee.retention_prob if employee.retention_prob is not None else 0.5
    )
    score = (1.0 - sentiment) * 0.6 + (1.0 - retention) * 0.4
    if employee.is_at_risk:
        score += 0.2
    return round(min(1.0, max(0.0, score)), 3)


def _top_match_score_candidate(candidate: CandidateTable) -> float:
    match_score = float(
        candidate.match_score if candidate.match_score is not None else 0.0
    )
    sentiment = float(candidate.sentiment_score or 0.0)
    score = (match_score * 0.8) + (sentiment * 0.2)
    return round(min(1.0, max(0.0, score)), 3)


def _workspace_snapshot(db: Session, user_text: str = "") -> Dict[str, Any]:
    query_lower = user_text.lower() if user_text else ""
    snapshot = {}

    # Check query intent keywords
    wants_risk = any(k in query_lower for k in ["risk", "morale", "sentiment", "unhappy", "leaving", "hazard", "at-risk", "employee"])
    wants_candidates = any(k in query_lower for k in ["candidate", "scout", "match", "hiring", "talent"])
    wants_integrations = any(k in query_lower for k in ["integration", "connection", "connect", "greenhouse", "slack", "jira", "workday"])
    wants_compliance = any(k in query_lower for k in ["compliance", "policy", "policies", "gate", "release"])
    wants_interventions = any(k in query_lower for k in ["intervention", "mitigate", "action"])
    wants_events = any(k in query_lower for k in ["event", "quarantine", "ops", "bronze", "silver", "sync"])

    # If it is a generic query and none of the above are matched, we don't fetch any snapshot details!
    if not (wants_risk or wants_candidates or wants_integrations or wants_compliance or wants_interventions or wants_events):
        return {}

    # Query only what is requested!
    if wants_risk:
        employees = filter_real_records(db.exec(select(EmployeeTable)).all())
        dept_summary = _department_risk_summary(employees)
        top_risk_employees = sorted(employees, key=_top_risk_score_employee, reverse=True)[:5]
        workforce = _compute_dashboard_snapshot(db)
        
        snapshot["workforce"] = {
            "total_workforce": workforce["total_workforce"],
            "at_risk": workforce["at_risk"],
            "at_risk_pct": (
                round((workforce["at_risk"] / workforce["total_workforce"]) * 100, 1)
                if workforce["total_workforce"]
                else 0.0
            ),
            "avg_morale": workforce["avg_morale"],
            "departments": workforce["departments"],
            "department_risk_clusters": dept_summary[:6],
            "top_risk_department": max(
                dept_summary, key=lambda item: item["at_risk_rate"], default=None
            ),
        }
        snapshot["directory"] = {
            "top_risk_employees": [
                {
                    "full_name": e.full_name,
                    "email": e.email,
                    "department": e.department,
                    "role": e.role,
                    "sentiment_score": e.sentiment_score,
                    "retention_prob": e.retention_prob,
                    "is_at_risk": e.is_at_risk,
                }
                for e in top_risk_employees
            ]
        }

    if wants_candidates:
        candidates = filter_real_records(db.exec(select(CandidateTable)).all())
        candidate_match_avg = (
            round(sum(float(c.match_score or 0.0) for c in candidates) / len(candidates), 3)
            if candidates
            else 0.0
        )
        top_candidate_matches = sorted(
            candidates, key=_top_match_score_candidate, reverse=True
        )[:5]
        
        snapshot["talent_scout"] = {
            "candidates_total": len(candidates),
            "candidate_match_avg": candidate_match_avg,
            "top_candidates": [
                {
                    "full_name": c.full_name,
                    "email": c.email,
                    "department": c.department,
                    "role": c.role,
                    "sentiment_score": c.sentiment_score,
                    "match_score": c.match_score,
                }
                for c in top_candidate_matches
            ]
        }

    if wants_integrations:
        integrations = db.exec(
            select(IntegrationConnectionTable).order_by(
                IntegrationConnectionTable.created_at.desc()
            )
        ).all()
        snapshot["enterprise_ops"] = {
            "integration_connections": [
                {
                    "name": c.name,
                    "provider": c.provider,
                    "source_type": c.source_type,
                    "status": c.status,
                }
                for c in integrations[:10]
            ]
        }

    if wants_compliance:
        policies = db.exec(
            select(CompliancePolicyTable)
            .where(CompliancePolicyTable.status == "active")
            .order_by(CompliancePolicyTable.created_at.desc())
        ).all()
        release_gates = db.exec(
            select(ReleaseGateTable).order_by(ReleaseGateTable.created_at.desc()).limit(5)
        ).all()
        snapshot["enterprise_ops"] = snapshot.get("enterprise_ops", {})
        snapshot["enterprise_ops"]["policy_packs"] = [
            {
                "policy_name": p.policy_name,
                "region": p.region,
                "action_type": p.action_type,
                "status": p.status,
            }
            for p in policies[:5]
        ]
        snapshot["enterprise_ops"]["release_gates"] = [
            {
                "artifact_name": g.artifact_name,
                "version": g.version,
                "status": g.status,
            }
            for g in release_gates
        ]

    if wants_interventions:
        interventions = db.exec(
            select(InterventionTable)
            .order_by(InterventionTable.created_at.desc())
            .limit(5)
        ).all()
        snapshot["enterprise_ops"] = snapshot.get("enterprise_ops", {})
        snapshot["enterprise_ops"]["active_interventions"] = [
            {
                "title": i.title,
                "priority": i.priority,
                "status": i.status,
                "target_department": i.target_department,
            }
            for i in interventions
        ]

    if wants_events:
        raw_events_count = len(db.exec(select(RawEventTable)).all())
        quarantine_count = len(db.exec(select(QuarantineEventTable)).all())
        connector_jobs = db.exec(
            select(ConnectorSyncJobTable)
            .order_by(ConnectorSyncJobTable.started_at.desc())
            .limit(5)
        ).all()
        snapshot["data_ops"] = {
            "raw_events": raw_events_count,
            "quarantine_events": quarantine_count,
            "sync_jobs": [
                {
                    "provider": j.provider,
                    "source_type": j.source_type,
                    "status": j.status,
                }
                for j in connector_jobs
            ]
        }

    return snapshot


def _is_workspace_query(user_text: str) -> bool:
    text = user_text.lower()
    terms = [
        "directory",
        "sentiment",
        "morale",
        "risk",
        "retention",
        "talent scout",
        "talent scout",
        "workflow",
        "workflows",
        "candidate",
        "workforce",
        "analytics",
        "intervention",
        "policy",
        "system status",
        "enterprise operations",
        "data ops",
        "model",
        "drift",
        "release gate",
        "runbook",
        "procurement",
        "integration",
        "quarantine",
        "department",
        "cluster",
        "priority",
        "predictive",
    ]
    return any(term in text for term in terms)


def _direct_workspace_answer(db: Session, user_text: str) -> str | None:
    if not _is_workspace_query(user_text):
        return None

    snapshot = _workspace_snapshot(db)
    workforce = snapshot["workforce"]
    directory = snapshot["directory"]
    talent = snapshot["talent_scout"]
    ops = snapshot["enterprise_ops"]
    data_ops = snapshot["data_ops"]

    lines = [
        "## Aurelinx Workspace Snapshot",
        f"- Workforce: **{workforce['total_workforce']}**",
        f"- At risk employees: **{workforce['at_risk']}**",
        f"- Average morale: **{workforce['avg_morale']}**",
        f"- Candidates: **{talent['candidates_total']}**",
        f"- Candidate match average: **{talent['candidate_match_avg']}**",
        f"- Raw events: **{data_ops['raw_events']}**",
        f"- Quarantine events: **{data_ops['quarantine_events']}**",
        f"- Active interventions: **{len(ops['active_interventions'])} shown**",
        f"- Active policy packs: **{len(ops['policy_packs'])} shown**",
    ]

    top_cluster = workforce.get("top_risk_department")
    if top_cluster:
        lines.append(
            f"- Top risk department: **{top_cluster['department']}** "
            f"({top_cluster['at_risk']} / {top_cluster['total']} at risk rate {top_cluster['at_risk_rate']:.3f})"
        )

    lines.append("")
    lines.append("### Risk Drivers")
    for item in workforce["top_risk_drivers"]:
        lines.append(f"- {item['factor']}: `{item['count']}`")

    if workforce.get("department_risk_clusters"):
        lines.append("")
        lines.append("### Department Risk Clusters")
        for cluster in workforce["department_risk_clusters"][:5]:
            lines.append(
                f"- {cluster['department']} | total `{cluster['total']}` | at risk `{cluster['at_risk']}` | "
                f"rate `{cluster['at_risk_rate']:.3f}`"
            )

    if directory["top_risk_employees"]:
        lines.append("")
        lines.append("### Highest Risk Employees")
        for e in directory["top_risk_employees"]:
            lines.append(
                f"- {e['full_name']} | {e['department']} | {e['role']} | "
                f"sentiment `{e['sentiment_score']}` | retention `{e['retention_prob']}` | risk `{e['risk_score']}`"
            )

    if talent["top_candidates"]:
        lines.append("")
        lines.append("### Top Candidate Matches")
        for c in talent["top_candidates"]:
            lines.append(
                f"- {c['full_name']} | {c['department']} | {c['role']} | "
                f"match `{c['match_score']}` | scout `{c['scout_score']}`"
            )

    lines.append("")
    lines.append("### Enterprise Operations")
    lines.append(f"- Integrations: `{len(ops['integration_connections'])}`")
    lines.append(f"- Model cards: `{len(ops['model_governance'])}`")
    lines.append(f"- Drift snapshots: `{len(ops['drift_snapshots'])}`")
    lines.append(f"- Release gates: `{len(ops['release_gates'])}`")
    lines.append(f"- DR runbooks: `{len(ops['dr_runbooks'])}`")
    lines.append(f"- Procurement artifacts: `{len(ops['procurement_artifacts'])}`")
    surface = snapshot.get("application_surface", {})
    if surface:
        lines.append("")
        lines.append("### Application Surfaces")
        lines.append(f"- Modules: {', '.join(surface.get('modules', []))}")
        lines.append(f"- Data planes: {', '.join(surface.get('data_planes', []))}")
        lines.append(f"- Import paths: {', '.join(surface.get('import_paths', []))}")

    return "\n".join(lines)


def _workspace_snapshot_summary_lines(snapshot: Dict[str, Any]) -> List[str]:
    workforce = snapshot.get("workforce", {})
    talent = snapshot.get("talent_scout", {})
    data_ops = snapshot.get("data_ops", {})
    ops = snapshot.get("enterprise_ops", {})
    top_cluster = workforce.get("top_risk_department")

    lines = [
        f"- Workforce: {workforce.get('total_workforce', 0)} total, {workforce.get('at_risk', 0)} at risk, avg morale {workforce.get('avg_morale', 0.0)}",
        f"- Candidate pool: {talent.get('candidates_total', 0)} total, avg match {talent.get('candidate_match_avg', 0.0)}",
        f"- Data ops: raw events {data_ops.get('raw_events', 0)}, quarantine {data_ops.get('quarantine_events', 0)}, canonical employees {data_ops.get('canonical_employees', 0)}, canonical candidates {data_ops.get('canonical_candidates', 0)}",
        f"- Enterprise ops: integrations {len(ops.get('integration_connections', []))}, interventions {len(ops.get('active_interventions', []))}, policies {len(ops.get('policy_packs', []))}",
        "- Application surfaces: dashboard, directory, sentiment, analytics, talent scout, workflow chat, intel center, enterprise ops, settings",
    ]
    if top_cluster:
        lines.append(
            f"- Top risk department: {top_cluster.get('department')} ({top_cluster.get('at_risk', 0)}/{top_cluster.get('total', 0)} at risk rate {top_cluster.get('at_risk_rate', 0.0):.3f})"
        )

    # 1. Top At-Risk Employees details
    top_risk_employees = snapshot.get("directory", {}).get("top_risk_employees", [])
    if top_risk_employees:
        lines.append("\nTop 5 At-Risk Employees:")
        for emp in top_risk_employees:
            lines.append(
                f"  * {emp.get('full_name')} ({emp.get('role')} - {emp.get('department')}): Sentiment {emp.get('sentiment_score')}, Retention Prob {emp.get('retention_prob')}, At Risk Flag: {emp.get('is_at_risk')}"
            )

    # 2. Top Candidate Matches details
    top_candidate_matches = talent.get("top_candidates", [])
    if top_candidate_matches:
        lines.append("\nTop 5 Candidate Matches:")
        for cand in top_candidate_matches:
            lines.append(
                f"  * {cand.get('full_name')} ({cand.get('role')} - {cand.get('department')}): Match Score {cand.get('match_score')}, Sentiment {cand.get('sentiment_score')}"
            )

    # 3. Active Connections details
    connections = ops.get("integration_connections", [])
    if connections:
        lines.append("\nIntegration Connections:")
        for conn in connections:
            lines.append(f"  * {conn.get('name')} ({conn.get('provider')} - {conn.get('source_type')}): Status {conn.get('status')}")

    # 4. Compliance Policies details
    policies = ops.get("policy_packs", [])
    if policies:
        lines.append("\nActive Compliance Policies:")
        for p in policies:
            lines.append(f"  * {p.get('policy_name')} (Region: {p.get('region')}, Action: {p.get('action_type')}): Status {p.get('status')}")

    # 5. Active Interventions details
    interventions = ops.get("active_interventions", [])
    if interventions:
        lines.append("\nActive Interventions:")
        for i in interventions:
            lines.append(f"  * {i.get('title')} (Priority: {i.get('priority')}): Status {i.get('status')}")

    return lines


def _record_counts(db: Session) -> Dict[str, int]:
    def _count(model) -> int:
        value = db.exec(select(func.count()).select_from(model)).one()
        if isinstance(value, tuple):
            value = value[0]
        return int(value or 0)

    return {
        "employees": _count(EmployeeTable),
        "candidates": _count(CandidateTable),
        "skills": _count(SkillTable),
        "experience": _count(ExperienceTable),
    }


def _is_count_query(user_text: str) -> bool:
    text = user_text.lower()
    count_phrases = (
        "how many",
        "number of",
        "count of",
        "total",
        "how much",
    )
    entity_phrases = (
        "candidate",
        "candidates",
        "employee",
        "employees",
        "workforce",
        "people",
    )
    return any(p in text for p in count_phrases) and any(
        p in text for p in entity_phrases
    )


def _direct_count_answer(db: Session, user_text: str) -> str | None:
    if not _is_count_query(user_text):
        return None

    counts = _record_counts(db)
    text = user_text.lower()
    if any(
        term in text
        for term in [
            "candidate",
            "candidates",
            "recruit",
            "hiring",
            "talent",
            "talents",
        ]
    ):
        return (
            f"There are **{counts['candidates']} candidates** in Postgres.\n\n"
            f"- Candidates: `{counts['candidates']}`\n"
            f"- Skills linked: `{counts['skills']}`\n"
            f"- Experience rows: `{counts['experience']}`"
        )
    if any(term in text for term in ["employee", "employees", "workforce", "people"]):
        return (
            f"There are **{counts['employees']} employees** in Postgres.\n\n"
            f"- Employees: `{counts['employees']}`\n"
            f"- Candidates: `{counts['candidates']}`"
        )
    return None


def _parse_csv_and_ingest(db: Session, attachments: List[ChatAttachmentTable]) -> Dict:
    import csv

    ingest_log = {
        "ingested": False,
        "employees": 0,
        "candidates": 0,
        "companies": 0,
        "errors": [],
        "actions": [],
    }
    for att in attachments:
        path = Path(att.file_path)
        if not path.exists() or path.suffix.lower() != ".csv":
            continue
        try:
            with open(path, mode="r", encoding="utf-8", errors="ignore") as f:
                content_sample = f.read(2048)
                f.seek(0)
                delimiter = ","
                if ";" in content_sample and "," not in content_sample:
                    delimiter = ";"

                reader = csv.DictReader(f, delimiter=delimiter)
                for row in reader:
                    if not row:
                        continue
                    clean_row = {
                        str(k).strip().lower(): str(v).strip()
                        for k, v in row.items()
                        if k is not None
                    }

                    email = (
                        clean_row.get("email")
                        or clean_row.get("email_address")
                        or clean_row.get("mail")
                    )
                    name = (
                        clean_row.get("full_name")
                        or clean_row.get("name")
                        or clean_row.get("employee_name")
                    )
                    role = (
                        clean_row.get("role")
                        or clean_row.get("position")
                        or clean_row.get("title")
                        or "Software Engineer"
                    )
                    dept = (
                        clean_row.get("department")
                        or clean_row.get("dept")
                        or "Engineering"
                    )

                    if not email or not name:
                        continue

                    is_candidate = (
                        "candidate" in clean_row
                        or clean_row.get("status", "").lower() == "candidate"
                        or clean_row.get("type", "").lower() == "candidate"
                    )

                    if is_candidate:
                        exists = db.exec(
                            select(CandidateTable).where(CandidateTable.email == email)
                        ).first()
                        if not exists:
                            cand = CandidateTable(
                                full_name=name,
                                email=email,
                                department=dept.title(),
                                role=role.title(),
                                sentiment_score=float(
                                    clean_row.get("sentiment_score")
                                    or clean_row.get("sentiment")
                                    or 0.65
                                ),
                            )
                            db.add(cand)
                            ingest_log["candidates"] += 1
                            ingest_log["actions"].append(
                                f"Ingested Candidate: {name} ({email})"
                            )
                    else:
                        exists = db.exec(
                            select(EmployeeTable).where(EmployeeTable.email == email)
                        ).first()
                        if not exists:
                            sentiment = 0.75
                            try:
                                sentiment = float(
                                    clean_row.get("sentiment_score")
                                    or clean_row.get("sentiment")
                                    or 0.75
                                )
                            except ValueError:
                                pass

                            is_at_risk = clean_row.get(
                                "is_at_risk", "false"
                            ).lower() in ["true", "1", "yes"]

                            retention = 0.90
                            try:
                                retention = float(
                                    clean_row.get("retention_prob")
                                    or clean_row.get("retention")
                                    or 0.90
                                )
                            except ValueError:
                                pass

                            emp = EmployeeTable(
                                full_name=name,
                                email=email,
                                department=dept.title(),
                                role=role.title(),
                                sentiment_score=sentiment,
                                is_at_risk=is_at_risk,
                                retention_prob=retention,
                            )
                            db.add(emp)
                            ingest_log["employees"] += 1
                            ingest_log["actions"].append(
                                f"Ingested Employee: {name} ({email})"
                            )

                    company = (
                        clean_row.get("company")
                        or clean_row.get("previous_company")
                        or clean_row.get("employer")
                    )
                    if company:
                        ingest_log["companies"] += 1

            db.commit()
            ingest_log["ingested"] = True
        except Exception as e:
            db.rollback()
            ingest_log["errors"].append(str(e))

    return ingest_log


def _apply_data_mutation(db: Session, user_text: str) -> Dict:
    """
    Highly powerful natural language data orchestration mutation layer.
    Supports creating, updating and fixing all entries:
    - set employee <email> risk true|false
    - move employee <email> to department <name>
    - add employee <name>, email <email>, role <role>, dept <dept>
    - add candidate <name>, email <email>, role <role>, dept <dept>
    - update employee <email> set <role|department|name> to <value>
    - add connection <name> provider <provider> type <type>
    """
    lowered = user_text.lower()
    mutation_log = {"updated": False, "actions": []}

    # Extract email and standard fields robustly
    email = None
    email_match = re.search(r"([a-zA-Z0-9\.\-\_\+]+@[a-zA-Z0-9\.\-\_]+\.[a-zA-Z]{2,})", user_text)
    if email_match:
        email = email_match.group(1)

    # 1. Check for risk updates
    if "risk" in lowered:
        val_match = re.search(r"\b(true|false|yes|no|high|low)\b", lowered)
        if email and val_match:
            value = val_match.group(1) in ["true", "yes", "high"]
            emp = db.exec(select(EmployeeTable).where(EmployeeTable.email == email)).first()
            if emp:
                emp.is_at_risk = value
                emp.updated_at = datetime.utcnow()
                db.add(emp)
                mutation_log["updated"] = True
                mutation_log["actions"].append(f"Set risk={value} for {email}")

    # 2. Check for department moves
    if any(k in lowered for k in ["move", "transfer", "change department", "dept"]):
        dept_match = re.search(r"(?:to|department|dept)\s+(?:is\s+|:\s*|of\s+)?([a-zA-Z0-9\s\-\_]+?)(?:,|$|\.|\s+for|\s+email)", user_text, re.IGNORECASE)
        if email and dept_match:
            dept = dept_match.group(1).strip().title()
            for prefix in ["department", "dept"]:
                if dept.lower().startswith(prefix + " "):
                    dept = dept[len(prefix)+1:].strip()
            emp = db.exec(select(EmployeeTable).where(EmployeeTable.email == email)).first()
            if emp:
                emp.department = dept
                emp.updated_at = datetime.utcnow()
                db.add(emp)
                mutation_log["updated"] = True
                mutation_log["actions"].append(f"Moved {email} to {dept}")

    # 3. Add Employee / Candidate
    is_add_employee = ("employee" in lowered) and any(kw in lowered for kw in ["add", "create", "insert", "register"])
    is_add_candidate = ("candidate" in lowered) and any(kw in lowered for kw in ["add", "create", "insert", "register"])

    if is_add_employee or is_add_candidate:
        name = None
        name_patterns = [
            r"(?:name|employee|candidate)\s+(?:is\s+|:\s*|named\s+)?([a-zA-Z\s]+?)(?:,|$|\s+email|\s+role|\s+dept|\s+department|\s+with)",
            r"(?:add|create|insert|register)\s+(?:employee|candidate|new\s+employee|new\s+candidate)?\s+([a-zA-Z\s]+?)(?:,|$|\s+email|\s+role|\s+dept|\s+department|\s+with)"
        ]
        for pattern in name_patterns:
            m = re.search(pattern, user_text, re.IGNORECASE)
            if m:
                name = m.group(1).strip()
                for prefix in ["named", "is", "employee", "candidate", "new", "a"]:
                    if name.lower().startswith(prefix + " "):
                        name = name[len(prefix)+1:].strip()
                break

        role = "Software Engineer"
        role_match = re.search(r"role\s+(?:is\s+|:\s*|to\s+)?([a-zA-Z0-9\s\-\_]+?)(?:,|$|\.|\s+email|\s+dept|\s+department|\s+with)", user_text, re.IGNORECASE)
        if role_match:
            role = role_match.group(1).strip().title()

        dept = "Engineering"
        dept_match = re.search(r"(?:dept|department)\s+(?:is\s+|:\s*|to\s+)?([a-zA-Z0-9\s\-\_]+?)(?:,|$|\.|\s+email|\s+role|\s+with)", user_text, re.IGNORECASE)
        if dept_match:
            dept = dept_match.group(1).strip().title()

        if name and email:
            if is_add_candidate:
                exists = db.exec(
                    select(CandidateTable).where(CandidateTable.email == email)
                ).first()
                if not exists:
                    cand = CandidateTable(
                        full_name=name,
                        email=email,
                        role=role,
                        department=dept,
                        sentiment_score=0.65,
                        match_score=0.85,
                    )
                    db.add(cand)
                    mutation_log["updated"] = True
                    mutation_log["actions"].append(
                        f"Created candidate {name} ({email}) in department {dept} as {role}"
                    )
            else:
                exists = db.exec(
                    select(EmployeeTable).where(EmployeeTable.email == email)
                ).first()
                if not exists:
                    emp = EmployeeTable(
                        full_name=name,
                        email=email,
                        role=role,
                        department=dept,
                        sentiment_score=0.75,
                        is_at_risk=False,
                        retention_prob=0.95,
                    )
                    db.add(emp)
                    mutation_log["updated"] = True
                    mutation_log["actions"].append(
                        f"Created employee {name} ({email}) in department {dept} as {role}"
                    )

    # 4. General updates
    if any(k in lowered for k in ["update", "correct", "change", "fix", "set"]) and email:
        emp = db.exec(select(EmployeeTable).where(EmployeeTable.email == email)).first()
        if emp:
            updated = False
            role_match = re.search(r"role\s+(?:to\s+|set\s+|is\s+|:\s*)?([a-zA-Z0-9\s\-\_]+?)(?:,|$|\.|\s+for)", user_text, re.IGNORECASE)
            if role_match:
                val = role_match.group(1).strip().title()
                emp.role = val
                mutation_log["actions"].append(f"Updated role to {val} for {email}")
                updated = True
            
            dept_match = re.search(r"(?:dept|department)\s+(?:to\s+|set\s+|is\s+|:\s*)?([a-zA-Z0-9\s\-\_]+?)(?:,|$|\.|\s+for)", user_text, re.IGNORECASE)
            if dept_match:
                val = dept_match.group(1).strip().title()
                emp.department = val
                mutation_log["actions"].append(f"Updated department to {val} for {email}")
                updated = True

            name_match = re.search(r"name\s+(?:to\s+|set\s+|is\s+|:\s*)?([a-zA-Z\s]+?)(?:,|$|\.|\s+for)", user_text, re.IGNORECASE)
            if name_match:
                val = name_match.group(1).strip().title()
                emp.full_name = val
                mutation_log["actions"].append(f"Updated full name to {val} for {email}")
                updated = True

            if updated:
                emp.updated_at = datetime.utcnow()
                db.add(emp)
                mutation_log["updated"] = True

    # 5. Integration connection mutations
    if (
        "add connection" in lowered
        or "create connection" in lowered
        or "add company" in lowered
    ):
        name_match = re.search(
            r"(?:connection|company)\s+([a-zA-Z0-9\s]+)(?:,|$)",
            user_text,
            re.IGNORECASE,
        )
        provider_match = re.search(
            r"provider\s+([a-zA-Z0-9\s]+)", user_text, re.IGNORECASE
        )
        type_match = re.search(r"type\s+([a-zA-Z0-9\s]+)", user_text, re.IGNORECASE)

        c_name = (
            name_match.group(1).strip() if name_match else "Greenhouse ATS Connection"
        )
        provider = (
            provider_match.group(1).strip().lower() if provider_match else "greenhouse"
        )
        source_type = type_match.group(1).strip().lower() if type_match else "ats"

        exists = db.exec(
            select(IntegrationConnectionTable).where(
                IntegrationConnectionTable.name == c_name
            )
        ).first()
        if not exists:
            conn = IntegrationConnectionTable(
                name=c_name, provider=provider, source_type=source_type, status="active"
            )
            db.add(conn)
            mutation_log["updated"] = True
            mutation_log["actions"].append(
                f"Configured integration connection {c_name} (provider: {provider})"
            )

    if mutation_log["updated"]:
        db.commit()

    return mutation_log


def _verify_mutation(db: Session, user_text: str, mutation_log: Dict[str, Any]) -> Dict[str, Any]:
    """Read back the exact target after a permitted mutation.

    This is intentionally a separate tool step. A successful write is not
    treated as complete until the committed row can be observed again.
    """
    if mutation_log.get("blocked"):
        return {"verified": False, "reason": "Mutation was blocked by policy"}
    if not mutation_log.get("updated"):
        return {
            "verified": False,
            "reason": "No database change was reported; nothing was verified",
            "actions": mutation_log.get("actions", []),
        }

    email_match = re.search(
        r"([a-zA-Z0-9.\-_+]+@[a-zA-Z0-9.\-_]+\.[a-zA-Z]{2,})", user_text
    )
    if not email_match:
        return {"verified": False, "reason": "Mutation result did not include a verifiable email target"}

    email = email_match.group(1)
    lowered = user_text.lower()
    is_candidate = "candidate" in lowered and "employee" not in lowered
    row = (
        db.exec(select(CandidateTable).where(CandidateTable.email == email)).first()
        if is_candidate
        else db.exec(select(EmployeeTable).where(EmployeeTable.email == email)).first()
    )
    if not row:
        return {"verified": False, "reason": f"Committed target {email} was not found on read-back"}

    return {
        "verified": True,
        "entity": "candidate" if is_candidate else "employee",
        "id": str(row.id),
        "email": row.email,
        "full_name": row.full_name,
        "role": row.role,
        "department": row.department,
    }


def _attachment_text_context(
    attachments: List[ChatAttachmentTable], max_chars: int = 12000
) -> str:
    chunks = []
    consumed = 0
    for att in attachments[:8]:
        path = Path(att.file_path)
        if not path.exists():
            continue
        suffix = path.suffix.lower()
        try:
            if suffix in [".txt", ".md", ".csv", ".json", ".log"]:
                text = path.read_text(encoding="utf-8", errors="ignore")
            elif suffix == ".pdf" and PdfReader is not None:
                reader = PdfReader(str(path))
                pages = []
                for page in reader.pages[:8]:
                    pages.append(page.extract_text() or "")
                text = "\n".join(pages)
            elif suffix == ".docx" and Document is not None:
                doc = Document(str(path))
                text = "\n".join([p.text for p in doc.paragraphs[:500]])
            elif (
                suffix in [".png", ".jpg", ".jpeg", ".webp"]
                and pytesseract is not None
                and Image is not None
            ):
                text = pytesseract.image_to_string(Image.open(str(path)))
            else:
                text = f"[Unsupported rich parsing for file type {suffix}]"
        except Exception:
            text = "[Failed to parse attachment]"
        block = f"\n[Attachment: {att.original_name}]\n{text[:2500]}\n"
        if consumed + len(block) > max_chars:
            break
        chunks.append(block)
        consumed += len(block)
    return "".join(chunks).strip()


def _parse_attachment_for_index(path: Path) -> Tuple[str, str, str]:
    suffix = path.suffix.lower()
    try:
        if suffix in [".txt", ".md", ".csv", ".json", ".log"]:
            return path.read_text(encoding="utf-8", errors="ignore"), "parsed", ""
        if suffix == ".pdf":
            if PdfReader is None:
                return "", "unsupported", "pypdf not installed"
            reader = PdfReader(str(path))
            pages = [(p.extract_text() or "") for p in reader.pages[:8]]
            return "\n".join(pages), "parsed", ""
        if suffix == ".docx":
            if Document is None:
                return "", "unsupported", "python-docx not installed"
            doc = Document(str(path))
            return "\n".join([p.text for p in doc.paragraphs[:500]]), "parsed", ""
        if suffix in [".png", ".jpg", ".jpeg", ".webp"]:
            if pytesseract is None or Image is None:
                return "", "unsupported", "ocr dependencies not installed"
            return pytesseract.image_to_string(Image.open(str(path))), "parsed", ""
        return "", "unsupported", f"unsupported file type: {suffix}"
    except Exception as e:
        return "", "failed", str(e)


async def _llm_stream_response(
    provider: str,
    api_key: str,
    base_url: str,
    model: str,
    user_text: str,
    context_payload: Dict,
):
    provider = (provider or "lmstudio").lower()
    casual_chat = _is_casual_chat(user_text)

    if provider == "lmstudio":
        endpoint = (
            f"{normalize_local_provider_base(base_url).rstrip('/')}/chat/completions"
        )
        model_name = model or "liquid/lfm2.5-1.2b"
        headers = {"Content-Type": "application/json"}
    elif provider == "opencode":
        endpoint = (
            f"{(base_url or 'https://opencode.ai/zen/v1').rstrip('/')}/chat/completions"
        )
        model_name = model or "gpt-5.5"
        headers = {"Content-Type": "application/json"}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
    elif provider == "openai":
        endpoint = "https://api.openai.com/v1/chat/completions"
        model_name = model or "gpt-4o-mini"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        }
    elif provider == "groq":
        endpoint = "https://api.groq.com/openai/v1/chat/completions"
        model_name = model or "llama-3.1-70b-versatile"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        }
    elif provider == "claude":
        endpoint = "https://api.anthropic.com/v1/messages"
        model_name = model or "claude-3-5-sonnet-20241022"
        headers = {
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
        }
    elif provider == "gemini":
        endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{model or 'gemini-1.5-pro'}:streamGenerateContent?key={api_key}"
        headers = {"Content-Type": "application/json"}
    else:
        yield f"Provider stream adapter for {provider} not supported."
        return

    system = (
        "You are Aurelinx, the authoritative AI intelligence of the Aurelinx Management OS — "
        "an executive-grade HR platform.\n"
        "You have live access to the HR employee directory, candidates database, analytics, "
        "compliance policies, integrations, enterprise ops, and agentic tools.\n"
        "A structured workspace snapshot may be attached; use it as the source of truth.\n"
        "IMPORTANT RULES:\n"
        "- Always respond in clear, beautiful Markdown with tables, bullet points, bold text where appropriate.\n"
        "- Never dump raw JSON or technical payloads to the user.\n"
        "- Never fabricate workspace data. Employee, candidate, metric, policy, and system facts must come only from verified live tool output or clearly identified conversation context.\n"
        "- You may answer general knowledge and conversational questions naturally, but do not present unverified or time-sensitive claims as live Aurelinx facts.\n"
        "- If evidence is missing, say that it is not verified and explain which live query or source is needed. Never fill gaps with plausible names, numbers, or actions.\n"
        "- Clearly distinguish what was found, what was inferred, and what is only a recommendation.\n"
        "- Never claim a mutation, deletion, search, or external action happened unless the tool result explicitly confirms it.\n"
        "- Be natural, direct, and executive in tone.\n"
        "- If asked about employees, present the data in a clean formatted table or list.\n"
        "- If deletion is requested, firmly state it requires human approval and cannot be automated."
    )

    tool_ctx = context_payload.get("tool_context", {})
    tool_runs = tool_ctx.get("tool_runs", [])
    has_tool_output = bool(tool_runs)

    tool_summary_lines = []
    for run in tool_runs:
        tool_name = run.get("tool", "unknown")
        if run.get("denied"):
            tool_summary_lines.append(
                f"- {tool_name}: ACCESS DENIED — {run.get('reason', '')}"
            )
        elif run.get("blocked"):
            tool_summary_lines.append(
                f"- {tool_name}: BLOCKED — {run.get('reason', '')}"
            )
        else:
            output = run.get("output", [])
            tool_summary_lines.append(f"- {tool_name}: {json.dumps(output)}")

    integrations = tool_ctx.get("integration_connections", [])
    compliance = tool_ctx.get("compliance_policies", [])
    mutations = tool_ctx.get("mutations", {})
    rbac = tool_ctx.get("rbac_role", "member")
    workspace_snapshot = tool_ctx.get("workspace_snapshot", {})

    if not casual_chat:
        tool_block = "\n".join(tool_summary_lines) if tool_summary_lines else "- No specific tool queries executed."
        workspace_block = (
            "\n".join(_workspace_snapshot_summary_lines(workspace_snapshot))
            if workspace_snapshot
            else ""
        )
        attachment_ctx = tool_ctx.get("attachment_text_context", "")
        user_content = (
            f"User request: {user_text}\n\n"
            f"--- WORKSPACE SNAPSHOT ---\n{workspace_block}\n"
            f"--- ATTACHMENT/FILE CONTENT ---\n{attachment_ctx}\n"
            f"--- LIVE TOOL DATA ---\n{tool_block}\n"
            f"Active integrations: {json.dumps(integrations)}\n"
            f"Compliance policies: {json.dumps(compliance)}\n"
            f"Mutation log: {json.dumps(mutations)}\n"
            f"User RBAC role: {rbac}\n"
            "--- END TOOL DATA ---\n\n"
            "Using the data above, answer the user's request clearly in Markdown. "
            "Present employee data as a formatted table. Do not include raw JSON in your response."
        )
    else:
        request_plan = context_payload.get("request_plan", {})
        user_content = (
            f"Request mode: {request_plan.get('mode', 'conversation')}\n"
            f"User request: {user_text}\n"
            "No live workspace retrieval was required for this request. Use the conversation history for continuity, answer naturally, and do not invent Aurelinx records or current metrics."
        )

    session_history = context_payload.get("session_history", [])
    history_messages = []
    for turn in session_history:
        role = turn.get("role", "user")
        content = turn.get("content", "").strip()
        if (
            not content
            or content.startswith("LLM failed")
            or content.startswith('{"tool_context')
        ):
            continue
        history_messages.append({"role": role, "content": content[:300]})

    if provider in ["openai", "lmstudio", "groq", "opencode"]:
        messages = [{"role": "system", "content": system}]
        messages.extend(history_messages[-6:])
        messages.append({"role": "user", "content": user_content})
        payload = {
            "model": model_name,
            "max_tokens": 1024,
            "temperature": 0.7 if casual_chat else 0.3,
            "messages": messages,
            "stream": True,
        }
    elif provider == "claude":
        messages = list(history_messages[-6:])
        messages.append({"role": "user", "content": user_content})
        payload = {
            "model": model_name,
            "max_tokens": 1024,
            "temperature": 0.7 if casual_chat else 0.3,
            "system": system,
            "messages": messages,
            "stream": True,
        }
    else:  # gemini
        full_prompt = f"{system}\n\n{user_content}"
        payload = {
            "contents": [{"parts": [{"text": full_prompt}]}],
            "generationConfig": {
                "temperature": 0.7 if casual_chat else 0.3,
                "maxOutputTokens": 1024,
            },
        }

    in_thinking = False
    async with httpx.AsyncClient(timeout=60.0) as client:
        async with client.stream(
            "POST", endpoint, json=payload, headers=headers
        ) as resp:
            resp.raise_for_status()

            if provider in ["openai", "lmstudio", "groq", "opencode"]:
                async for line in resp.aiter_lines():
                    if line.startswith("data:"):
                        data_str = line[5:].strip()
                        if data_str == "[DONE]":
                            break
                        try:
                            chunk = json.loads(data_str)
                            delta = chunk.get("choices", [{}])[0].get("delta", {})

                            content = delta.get("content")
                            reasoning = delta.get("reasoning_content")

                            if reasoning:
                                if not in_thinking:
                                    yield "<think>"
                                    in_thinking = True
                                yield reasoning
                            elif content:
                                if in_thinking:
                                    yield "</think>"
                                    in_thinking = False
                                yield content
                        except Exception:
                            continue

            elif provider == "claude":
                async for line in resp.aiter_lines():
                    if line.startswith("data:"):
                        data_str = line[5:].strip()
                        try:
                            chunk = json.loads(data_str)
                            if chunk.get("type") == "content_block_delta":
                                delta = chunk.get("delta", {})
                                text = delta.get("text", "")
                                if text:
                                    yield text
                        except Exception:
                            continue

            elif provider == "gemini":
                # Handle Gemini chunk-by-chunk stream response beautifully
                async for chunk_bytes in resp.aiter_bytes():
                    chunk_str = chunk_bytes.decode("utf-8", errors="ignore")
                    try:
                        obj = json.loads(chunk_str.strip())
                        text = (
                            obj.get("candidates", [{}])[0]
                            .get("content", {})
                            .get("parts", [{}])[0]
                            .get("text", "")
                        )
                        if text:
                            yield text
                    except Exception:
                        import re

                        for m in re.finditer(
                            r'"text"\s*:\s*"([^"\\]*(?:\\.[^"\\]*)*)"', chunk_str
                        ):
                            try:
                                yield m.group(1).encode().decode("unicode-escape")
                            except Exception:
                                pass

    if in_thinking:
        yield "</think>"


async def _llm_response(
    provider: str,
    api_key: str,
    base_url: str,
    model: str,
    user_text: str,
    context_payload: Dict,
    system_override: Optional[str] = None,
    temperature_override: Optional[float] = None,
) -> str:
    provider = (provider or "lmstudio").lower()

    if provider == "lmstudio":
        endpoint = (
            f"{normalize_local_provider_base(base_url).rstrip('/')}/chat/completions"
        )
        model_name = model or "liquid/lfm2.5-1.2b"
        headers = {"Content-Type": "application/json"}
    elif provider == "opencode":
        endpoint = (
            f"{(base_url or 'https://opencode.ai/zen/v1').rstrip('/')}/chat/completions"
        )
        model_name = model or "gpt-5.5"
        headers = {"Content-Type": "application/json"}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"
    elif provider == "openai":
        endpoint = "https://api.openai.com/v1/chat/completions"
        model_name = model or "gpt-4o-mini"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        }
    elif provider == "groq":
        endpoint = "https://api.groq.com/openai/v1/chat/completions"
        model_name = model or "llama-3.1-70b-versatile"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        }
    elif provider == "claude":
        endpoint = "https://api.anthropic.com/v1/messages"
        model_name = model or "claude-3-5-sonnet-20241022"
        headers = {
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
        }
    elif provider == "gemini":
        endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{model or 'gemini-1.5-pro'}:generateContent?key={api_key}"
        headers = {"Content-Type": "application/json"}
    else:
        return "Provider adapter for this chat request is not configured. Switch to LM Studio/OpenAI."

    casual_chat = _is_casual_chat(user_text)

    default_system = (
        "You are Aurelinx, the authoritative AI intelligence of the Aurelinx Management OS — "
        "an executive-grade HR platform.\n"
        "You have live access to the HR employee directory, candidates database, analytics, "
        "compliance policies, integrations, enterprise ops, and agentic tools.\n"
        "A structured workspace snapshot may be attached; use it as the source of truth.\n"
        "IMPORTANT RULES:\n"
        "- Always respond in clear, beautiful Markdown with tables, bullet points, bold text where appropriate.\n"
        "- Never dump raw JSON or technical payloads to the user.\n"
        "- Never fabricate workspace data. Employee, candidate, metric, policy, and system facts must come only from verified live tool output or clearly identified conversation context.\n"
        "- You may answer general knowledge and conversational questions naturally, but do not present unverified or time-sensitive claims as live Aurelinx facts.\n"
        "- If evidence is missing, say that it is not verified and explain which live query or source is needed. Never fill gaps with plausible names, numbers, or actions.\n"
        "- Clearly distinguish what was found, what was inferred, and what is only a recommendation.\n"
        "- Never claim a mutation, deletion, search, or external action happened unless the tool result explicitly confirms it.\n"
        "- Be natural, direct, and executive in tone.\n"
        "- If asked about employees, present the data in a clean formatted table or list.\n"
        "- If deletion is requested, firmly state it requires human approval and cannot be automated."
    )
    system = system_override or default_system
    response_temperature = (
        temperature_override
        if temperature_override is not None
        else (0.7 if casual_chat else 0.3)
    )

    # Build clean tool summary — never pass nested session_history inside the tool block
    tool_ctx = context_payload.get("tool_context", {})
    tool_runs = tool_ctx.get("tool_runs", [])
    has_tool_output = bool(tool_runs)

    # Produce a compact, readable tool summary string (not raw JSON of entire payload)
    tool_summary_lines = []
    for run in tool_runs:
        tool_name = run.get("tool", "unknown")
        if run.get("denied"):
            tool_summary_lines.append(
                f"- {tool_name}: ACCESS DENIED — {run.get('reason', '')}"
            )
        elif run.get("blocked"):
            tool_summary_lines.append(
                f"- {tool_name}: BLOCKED — {run.get('reason', '')}"
            )
        else:
            output = run.get("output", [])
            tool_summary_lines.append(f"- {tool_name}: {json.dumps(output)}")

    integrations = tool_ctx.get("integration_connections", [])
    compliance = tool_ctx.get("compliance_policies", [])
    mutations = tool_ctx.get("mutations", {})
    rbac = tool_ctx.get("rbac_role", "member")
    workspace_snapshot = tool_ctx.get("workspace_snapshot", {})

    if not casual_chat:
        tool_block = "\n".join(tool_summary_lines) if tool_summary_lines else "- No specific tool queries executed."
        workspace_block = (
            "\n".join(_workspace_snapshot_summary_lines(workspace_snapshot))
            if workspace_snapshot
            else ""
        )
        attachment_ctx = tool_ctx.get("attachment_text_context", "")
        user_content = (
            f"User request: {user_text}\n\n"
            f"--- WORKSPACE SNAPSHOT ---\n{workspace_block}\n"
            f"--- ATTACHMENT/FILE CONTENT ---\n{attachment_ctx}\n"
            f"--- LIVE TOOL DATA ---\n{tool_block}\n"
            f"Active integrations: {json.dumps(integrations)}\n"
            f"Compliance policies: {json.dumps(compliance)}\n"
            f"Mutation log: {json.dumps(mutations)}\n"
            f"User RBAC role: {rbac}\n"
            "--- END TOOL DATA ---\n\n"
            "Using the data above, answer the user's request clearly in Markdown. "
            "Present employee data as a formatted table. Do not include raw JSON in your response."
        )
    else:
        request_plan = context_payload.get("request_plan", {})
        user_content = (
            f"Request mode: {request_plan.get('mode', 'conversation')}\n"
            f"User request: {user_text}\n"
            "No live workspace retrieval was required for this request. Use the conversation history for continuity, answer naturally, and do not invent Aurelinx records or current metrics."
        )

    # Build multi-turn chat history from clean session context (never include tool dump messages)
    session_history = context_payload.get("session_history", [])
    history_messages = []
    for turn in session_history:
        role = turn.get("role", "user")
        content = turn.get("content", "").strip()
        # Skip any poisoned turns that contain raw JSON dumps or LLM failure messages
        if (
            not content
            or content.startswith("LLM failed")
            or content.startswith('{"tool_context')
        ):
            continue
        # Only keep clean conversational turns (cap at 300 chars per turn to stay token-safe)
        history_messages.append({"role": role, "content": content[:300]})

    if provider in ["openai", "lmstudio", "groq", "opencode"]:
        messages = [{"role": "system", "content": system}]
        messages.extend(history_messages[-6:])  # last 3 turns max
        messages.append({"role": "user", "content": user_content})
        payload = {
            "model": model_name,
            "max_tokens": 1024,
            "temperature": response_temperature,
            "messages": messages,
        }
    elif provider == "claude":
        messages = list(history_messages[-6:])
        messages.append({"role": "user", "content": user_content})
        payload = {
            "model": model_name,
            "max_tokens": 1024,
            "temperature": response_temperature,
            "system": system,
            "messages": messages,
        }
    else:  # gemini
        full_prompt = f"{system}\n\n{user_content}"
        payload = {
            "contents": [{"parts": [{"text": full_prompt}]}],
            "generationConfig": {
                "temperature": response_temperature,
                "maxOutputTokens": 1024,
            },
        }

    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.post(endpoint, json=payload, headers=headers)
        resp.raise_for_status()
        data = resp.json()
        if provider in ["openai", "lmstudio", "groq", "opencode"]:
            return _sanitize_llm_response(
                data["choices"][0]["message"]["content"], user_text
            )
        if provider == "claude":
            text_blocks = [
                b.get("text", "")
                for b in data.get("content", [])
                if isinstance(b, dict)
            ]
            return _sanitize_llm_response(
                "\n".join([t for t in text_blocks if t]).strip(), user_text
            )
        # gemini
        return _sanitize_llm_response(
            data.get("candidates", [{}])[0]
            .get("content", {})
            .get("parts", [{}])[0]
            .get("text", "")
            .strip(),
            user_text,
        )


def _request_plan(user_text: str, has_attachments: bool = False) -> Dict[str, Any]:
    """Route a request before retrieval or mutation.

    This is deliberately an operational intent router, not hidden model reasoning.
    It keeps general conversation useful while ensuring workspace facts trigger live
    tools and mutations/deletions are never inferred from a vague sentence.
    """
    text = user_text.strip().lower()
    mutation_words = ("add", "create", "insert", "update", "set", "move", "transfer", "correct", "fix", "change", "ingest", "import", "register")
    delete_words = ("delete", "remove", "purge", "clear")
    employee_words = ("employee", "employees", "staff", "workforce", "team member", "people", "directory", "worker")
    candidate_words = ("candidate", "candidates", "applicant", "applicants", "talent scout", "hiring", "recruit")
    analytics_words = ("dashboard", "analytics", "metric", "metrics", "sentiment", "morale", "retention", "at-risk", "at risk", "risk", "burnout")
    enterprise_words = ("integration", "connection", "compliance", "policy", "policies", "release gate", "quarantine", "sync", "runbook", "intervention")
    data_verbs = ("show", "find", "search", "list", "lookup", "look up", "who", "which", "count", "how many", "analyze", "compare", "rank", "top", "current", "live", "our", "my", "data", "records", "report")

    is_mutation = any(word in text for word in mutation_words)
    is_delete = any(word in text for word in delete_words)
    mentions_employee = any(word in text for word in employee_words)
    mentions_candidate = any(word in text for word in candidate_words)
    mentions_analytics = any(word in text for word in analytics_words)
    mentions_enterprise = any(word in text for word in enterprise_words)
    asks_for_data = any(word in text for word in data_verbs)
    asks_count = any(word in text for word in ("count", "how many", "total", "number of"))
    workspace_signal = (
        mentions_employee
        or mentions_candidate
        or mentions_analytics
        or mentions_enterprise
    ) and (asks_for_data or is_mutation or is_delete or len(text.split()) > 2)
    workspace_signal = workspace_signal or is_mutation or is_delete or has_attachments

    capabilities = ["conversation_context"]
    if workspace_signal:
        if mentions_employee or (not mentions_candidate and not mentions_enterprise and (mentions_analytics or asks_for_data)):
            capabilities.append("search_employees")
        if mentions_candidate:
            capabilities.append("search_candidates")
        if mentions_analytics or asks_count or (workspace_signal and not mentions_employee and not mentions_candidate and not mentions_enterprise):
            capabilities.append("dashboard_snapshot")
        if mentions_enterprise:
            capabilities.append("workspace_snapshot")
    if is_mutation:
        capabilities.append("mutate_data")
    if is_delete:
        capabilities.append("human_approval_delete")
    if has_attachments or any(word in text for word in ("csv", "excel", "spreadsheet", "file", "attachment")):
        capabilities.append("parse_csv_attachment")

    if is_delete:
        mode = "governed_mutation"
    elif is_mutation:
        mode = "mutation"
    elif workspace_signal:
        mode = "workspace_analysis"
    else:
        mode = "conversation"

    return {
        "mode": mode,
        "needs_live_data": workspace_signal,
        "capabilities": list(dict.fromkeys(capabilities)),
        "reason": "live workspace evidence requested" if workspace_signal else "general conversation or knowledge request",
    }


def _tool_policy(user_text: str, has_attachments: bool = False) -> List[str]:
    return _request_plan(user_text, has_attachments).get("capabilities", ["conversation_context"])


TOOL_RBAC = {
    "search_employees": ["member", "admin"],
    "search_candidates": ["member", "admin"],
    "dashboard_snapshot": ["member", "admin"],
    "mutate_data": ["admin"],
    "human_approval_delete": ["member", "admin"],
    "parse_csv_attachment": ["member", "admin"],
}


def _user_role(current_user: TokenData) -> str:
    return "admin" if current_user.is_admin else "member"


def _is_casual_chat(user_text: str) -> bool:
    return _request_plan(user_text).get("mode") == "conversation"


AGENT_TOOL_CATALOG = {
    "employee.search": "Search verified employee records using a natural-language query.",
    "candidate.search": "Search verified candidate records using a natural-language query.",
    "dashboard.snapshot": "Calculate current workforce totals, risk, morale, and department metrics.",
    "workspace.snapshot": "Retrieve the relevant verified workspace section for analytics, integrations, policy, or operations.",
    "document.csv_ingest": "Parse and validate an attached CSV file; only use when an attachment is present.",
    "data.mutate": "Create or update one explicitly requested employee, candidate, or integration record. Admin only.",
    "data.verify": "Read back a previously committed mutation to verify the exact database state.",
}


def _agent_tool_label(tool_name: str) -> str:
    return {
        "employee.search": "employee directory",
        "candidate.search": "candidate directory",
        "dashboard.snapshot": "workforce analytics",
        "workspace.snapshot": "workspace records",
        "document.csv_ingest": "the attached CSV file",
        "data.mutate": "the requested data change",
        "data.verify": "the saved record",
    }.get(tool_name, tool_name.replace(".", " "))


def _agent_tool_result_message(tool_name: str, result: Dict[str, Any]) -> str:
    if result.get("blocked"):
        return f"I could not run {_agent_tool_label(tool_name)} because policy blocked it."
    if tool_name in {"employee.search", "candidate.search"}:
        return f"I found {len(result.get('matches', []))} matching record(s) in {_agent_tool_label(tool_name)}."
    if tool_name == "data.mutate":
        mutation = result.get("result") or {}
        if mutation.get("updated"):
            return "The requested data change was committed. I will read it back before finishing."
        if mutation.get("actions"):
            return "The data-change step completed with policy results that need to be reviewed."
    if tool_name == "data.verify":
        verification = result.get("result") or {}
        return (
            "I read the saved record back and verification passed."
            if verification.get("verified")
            else "I read the saved record back, but verification did not pass."
        )
    if tool_name == "dashboard.snapshot":
        return "I calculated the current workforce analytics snapshot."
    if tool_name == "workspace.snapshot":
        return "I loaded the relevant workspace records."
    if tool_name == "document.csv_ingest":
        return "I parsed and validated the attached CSV file."
    return f"The {_agent_tool_label(tool_name)} step completed."


def _parse_agent_decision(raw_text: str) -> Dict[str, Any]:
    """Parse and strictly validate one model controller decision."""
    text = (raw_text or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*|\s*```$", "", text, flags=re.IGNORECASE | re.DOTALL).strip()
    try:
        payload = json.loads(text)
    except Exception:
        start = text.find("{")
        end = text.rfind("}")
        if start < 0 or end <= start:
            return {"action": "invalid", "error": "Controller did not return JSON"}
        try:
            payload = json.loads(text[start : end + 1])
        except Exception:
            return {"action": "invalid", "error": "Controller returned malformed JSON"}

    if not isinstance(payload, dict):
        return {"action": "invalid", "error": "Controller decision was not an object"}
    action = str(payload.get("action", "")).lower().strip()
    if action not in {"tool", "respond", "approval_required"}:
        return {"action": "invalid", "error": "Unsupported controller action"}
    if action == "tool":
        tool = str(payload.get("tool", "")).strip()
        if tool not in AGENT_TOOL_CATALOG:
            return {"action": "invalid", "error": f"Unsupported tool: {tool}"}
        args = payload.get("arguments")
        if not isinstance(args, dict):
            return {"action": "invalid", "error": "Tool arguments must be an object"}
        return {
            "action": "tool",
            "tool": tool,
            "arguments": args,
            "message": str(payload.get("message", "Model selected the next bounded tool"))[:500],
        }
    if action == "approval_required":
        return {
            "action": "approval_required",
            "message": str(payload.get("message", "Human approval is required before this action"))[:500],
        }
    return {
        "action": "respond",
        "message": str(payload.get("message", "Model selected a final response"))[:500],
        "answer": str(payload.get("answer", "")).strip()[:12000],
    }


def _agent_controller_prompt(
    user_text: str,
    history: List[Dict[str, str]],
    tool_transcript: List[Dict[str, Any]],
    current_user: TokenData,
    has_attachments: bool,
) -> str:
    catalog = "\n".join(f"- {name}: {description}" for name, description in AGENT_TOOL_CATALOG.items())
    return (
        "You are the Aurelinx execution controller. Choose exactly one next action and return ONLY valid JSON.\n"
        "You are not the private chain-of-thought writer. Do not expose chain-of-thought.\n"
        "Allowed JSON shapes:\n"
        '{"action":"tool","tool":"employee.search","arguments":{"query":"..."},"message":"short natural progress sentence shown in the activity timeline"}\n'
        '{"action":"respond","message":"short natural progress sentence shown before the answer","answer":"final user-facing answer for a conversation-only request"}\n'
        '{"action":"approval_required","message":"explain why admin approval is required"}\n'
        "Rules:\n"
        "- Select at most one tool per turn, then wait for its result.\n"
        "- Use live tools for Aurelinx workspace facts; do not guess records or metrics.\n"
        "- Never select data.mutate for a delete/remove request; select approval_required instead.\n"
        f"- Current user role: {'admin' if current_user.is_admin else 'member'}.\n"
        f"- Attached files available: {has_attachments}.\n"
        f"Available tools:\n{catalog}\n\n"
        f"Original user request:\n{user_text}\n\n"
        f"Conversation history:\n{json.dumps(history[-6:], default=str)}\n\n"
        f"Tool calls and results already observed:\n{json.dumps(tool_transcript[-8:], default=str)}\n\n"
        "The message field must be a natural operational update, not JSON, not a greeting, and not hidden chain-of-thought. Return one JSON decision now."
    )


def _execute_agent_tool(
    tool_name: str,
    arguments: Dict[str, Any],
    user_text: str,
    current_user: TokenData,
    session_id: str,
    mutation_state: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Execute exactly one allowlisted tool in an isolated DB session."""
    with Session(engine) as db:
        attachments = db.exec(
            select(ChatAttachmentTable).where(ChatAttachmentTable.session_id == str(session_id))
        ).all()
        query = str(arguments.get("query") or user_text)[:12000]
        if tool_name == "employee.search":
            rows = _search_employees(db, query, limit=min(int(arguments.get("limit", 8)), 20))
            return {
                "tool": tool_name,
                "matches": [
                    {"name": row.full_name, "email": row.email, "role": row.role, "department": row.department, "risk": row.is_at_risk}
                    for row in rows
                ],
            }
        if tool_name == "candidate.search":
            rows = _search_candidates(db, query, limit=min(int(arguments.get("limit", 8)), 20))
            return {
                "tool": tool_name,
                "matches": [
                    {"name": row.full_name, "email": row.email, "role": row.role, "department": row.department, "match_score": row.match_score}
                    for row in rows
                ],
            }
        if tool_name == "dashboard.snapshot":
            return {"tool": tool_name, "snapshot": _compute_dashboard_snapshot(db)}
        if tool_name == "workspace.snapshot":
            return {"tool": tool_name, "snapshot": _workspace_snapshot(db, query), "record_counts": _record_counts(db)}
        if tool_name == "document.csv_ingest":
            return {"tool": tool_name, "result": _parse_csv_and_ingest(db, attachments)}
        if tool_name == "data.mutate":
            if not current_user.is_admin:
                return {"tool": tool_name, "blocked": True, "reason": "RBAC requires an admin"}
            if any(word in user_text.lower() for word in ("delete", "remove", "purge", "clear")):
                return {"tool": tool_name, "blocked": True, "reason": "Deletion requires human approval"}
            return {"tool": tool_name, "result": _apply_data_mutation(db, user_text)}
        if tool_name == "data.verify":
            return {"tool": tool_name, "result": _verify_mutation(db, user_text, mutation_state or {})}
        return {"tool": tool_name, "blocked": True, "reason": "Tool was not available"}


def _sanitize_llm_response(text: str, user_text: str) -> str:
    """
    Ensure raw generated markdown text is passed directly to allow comprehensive markdown styling.
    """
    if not text:
        return ""
    return text.strip()


def _direct_casual_reply(user_text: str) -> str:
    # Retained as a fallback if model calls fail entirely
    return "Hi, I am Aurelinx. Let me know how I can assist with HR analytics, employee tracking, or database management."


def _provider_error_code(error: Any) -> str:
    text = str(error or "").lower()
    if "429" in text or "rate limit" in text or "too many requests" in text:
        return "MODEL_RATE_LIMITED"
    if "401" in text or "403" in text or "unauthorized" in text:
        return "MODEL_AUTH_FAILED"
    return "MODEL_PROVIDER_FAILED"


def _provider_error_label(error: Any) -> str:
    code = _provider_error_code(error)
    return {
        "MODEL_RATE_LIMITED": "provider rate limit (HTTP 429)",
        "MODEL_AUTH_FAILED": "provider authentication/authorization failure",
        "MODEL_PROVIDER_FAILED": "provider request failure",
    }.get(code, "provider request failure")


def _safe_provider_failure_reply(context_payload: Dict, error: Any = None) -> str:
    """Return a useful failure response without leaking raw internal payloads."""
    tool_context = context_payload.get("tool_context", {})
    runs = tool_context.get("tool_runs", [])
    summaries = []
    for run in runs:
        name = run.get("tool", "workflow tool")
        if run.get("denied") or run.get("blocked"):
            summaries.append(f"{name}: blocked by policy")
        else:
            output = run.get("output")
            if isinstance(output, list):
                summaries.append(f"{name}: {len(output)} records returned")
            elif isinstance(output, dict):
                summaries.append(f"{name}: completed")
            else:
                summaries.append(f"{name}: completed")
    evidence = "; ".join(summaries) if summaries else "no retrieval tools were required"
    reason = _provider_error_label(error) if error else "provider did not return a response"
    return (
        "I completed the permitted retrieval and safety checks, but the configured language "
        f"model provider returned a {reason}. No unsupported database change was made. "
        f"Operational results: {evidence}. Please retry or switch to another configured provider."
    )


def _build_context_payload(
    db: Session,
    session_id: UUID,
    user_text: str,
    current_user: TokenData,
    attachments: List[ChatAttachmentTable],
    workflow_run_id: Optional[str] = None,
    event_sink=None,
) -> Dict:
    """
    Build a lean, token-safe context payload.
    - Casual prompts skip history entirely.
    - Tool context is always included but session history is strictly filtered.
    - Poisoned turns (raw JSON dumps, LLM failure messages) are stripped at source.
    """
    request_plan = _request_plan(user_text, bool(attachments))
    tool_context = _execute_tools(
        db,
        user_text,
        current_user,
        attachments,
        workflow_run_id=workflow_run_id,
        event_sink=event_sink,
    )

    # Every request gets conversation context, including greetings. This is an
    # operational context tool, not a claim that the database was queried.
    history_rows = db.exec(
        select(ChatMessageTable)
        .where(ChatMessageTable.session_id == str(session_id))
        .order_by(ChatMessageTable.created_at.desc())
        .limit(6)
    ).all()

    # Strict filtering: drop poisoned/raw-JSON messages from history
    clean_history = []
    for m in reversed(history_rows):
        content = (m.content or "").strip()
        # Skip LLM failure fallback dumps and raw JSON tool context leaks
        if (
            content.startswith("LLM failed")
            or content.startswith('{"tool_context')
            or not content
        ):
            continue
        clean_history.append({"role": m.role, "content": content[:400]})

    if workflow_run_id:
        context_scope = ToolEventScope(
            workflow_run_id,
            "conversation.context",
            "context",
            "Loading relevant conversation context",
            sink=event_sink,
        )
        context_scope.start({"session_id": str(session_id), "turn_limit": 6})
        context_scope.complete({"messages_loaded": len(clean_history), "mode": request_plan["mode"]})
        tool_context.setdefault("tool_runs", []).append(
            {
                "tool": "conversation_context",
                "output": {"messages_loaded": len(clean_history), "mode": request_plan["mode"]},
            }
        )

    return {
        "tool_context": tool_context,
        "session_history": clean_history,
        "request_plan": request_plan,
    }


def _execute_tools(
    db: Session,
    user_text: str,
    current_user: TokenData,
    attachments: List[ChatAttachmentTable],
    workflow_run_id: Optional[str] = None,
    event_sink=None,
) -> Dict:
    request_plan = _request_plan(user_text, bool(attachments))
    tools = request_plan["capabilities"]
    result = {"tool_policy": tools, "tool_runs": [], "request_plan": request_plan}
    mutation_log = {"updated": False, "actions": [], "blocked": False}

    def scope(tool_name: str, phase: str, message: str) -> Optional[ToolEventScope]:
        if not workflow_run_id:
            return None
        current = ToolEventScope(
            workflow_run_id,
            tool_name,
            phase,
            message,
            sink=event_sink,
        )
        current.start({"query": user_text[:240]})
        return current

    role = _user_role(current_user)
    if "search_employees" in tools:
        tool_scope = scope(
            "employee.search",
            "retrieval",
            "Searching employee records",
        )
        if role in TOOL_RBAC["search_employees"]:
            emp_results = _search_employees(db, user_text, limit=8)
            result["tool_runs"].append(
                {
                    "tool": "search_employees",
                    "output": [
                        {
                            "name": e.full_name,
                            "email": e.email,
                            "role": e.role,
                            "department": e.department,
                            "risk": e.is_at_risk,
                        }
                        for e in emp_results
                    ],
                }
            )
            if tool_scope:
                tool_scope.complete({"matches_found": len(emp_results)})
        else:
            result["tool_runs"].append(
                {
                    "tool": "search_employees",
                    "denied": True,
                    "reason": "RBAC policy denied",
                }
            )
            if tool_scope:
                tool_scope.complete({"reason": "RBAC policy denied"}, status="blocked", error_code="RBAC_DENIED")
    if "search_candidates" in tools:
        tool_scope = scope(
            "candidate.search",
            "retrieval",
            "Searching candidate records",
        )
        if role in TOOL_RBAC["search_candidates"]:
            cand_results = _search_candidates(db, user_text, limit=8)
            result["tool_runs"].append(
                {
                    "tool": "search_candidates",
                    "output": [
                        {
                            "name": c.full_name,
                            "email": c.email,
                            "role": c.role,
                            "department": c.department,
                        }
                        for c in cand_results
                    ],
                }
            )
            if tool_scope:
                tool_scope.complete({"matches_found": len(cand_results)})
        else:
            result["tool_runs"].append(
                {
                    "tool": "search_candidates",
                    "denied": True,
                    "reason": "RBAC policy denied",
                }
            )
            if tool_scope:
                tool_scope.complete({"reason": "RBAC policy denied"}, status="blocked", error_code="RBAC_DENIED")
    if "dashboard_snapshot" in tools:
        tool_scope = scope(
            "dashboard.snapshot",
            "retrieval",
            "Calculating workforce dashboard snapshot",
        )
        if role in TOOL_RBAC["dashboard_snapshot"]:
            snapshot = _compute_dashboard_snapshot(db)
            result["tool_runs"].append(
                {
                    "tool": "dashboard_snapshot",
                    "output": snapshot,
                }
            )
            if tool_scope:
                tool_scope.complete({"record_count": snapshot.get("total_workforce", 0)})
        else:
            result["tool_runs"].append(
                {
                    "tool": "dashboard_snapshot",
                    "denied": True,
                    "reason": "RBAC policy denied",
                }
            )
            if tool_scope:
                tool_scope.complete({"reason": "RBAC policy denied"}, status="blocked", error_code="RBAC_DENIED")
    if "mutate_data" in tools:
        tool_scope = scope(
            "data.mutate",
            "mutation",
            "Preparing authorized database changes",
        )
        if role not in TOOL_RBAC["mutate_data"]:
            mutation_log["blocked"] = True
            mutation_log["actions"].append(
                "Mutation blocked: RBAC policy requires admin."
            )
        else:
            mutation_log = _apply_data_mutation(db, user_text)
        result["tool_runs"].append({"tool": "mutate_data", "output": mutation_log})
        if tool_scope:
            tool_scope.complete(
                {
                    "updated": mutation_log.get("updated", False),
                    "blocked": mutation_log.get("blocked", False),
                    "actions_count": len(mutation_log.get("actions", [])),
                },
                status="blocked" if mutation_log.get("blocked") else "completed",
                error_code="RBAC_DENIED" if mutation_log.get("blocked") else None,
            )
    if "human_approval_delete" in tools:
        tool_scope = scope(
            "approval.delete",
            "governance",
            "Blocking deletion pending human approval",
        )
        result["tool_runs"].append(
            {
                "tool": "human_approval_delete",
                "blocked": True,
                "reason": (
                    "Aurelinx Governance Protocol Violation: Delete actions cannot be automated by the AI agent "
                    "under any circumstances without manual Human-in-the-Loop approval. Safe abort triggered. "
                    "Instruct the user that manual confirmation is strictly required to delete this resource."
                ),
            }
        )
        if tool_scope:
            tool_scope.complete(
                {"approval_required": True},
                status="blocked",
                error_code="HUMAN_APPROVAL_REQUIRED",
            )
    if "parse_csv_attachment" in tools or any(att.original_name.lower().endswith(".csv") for att in attachments):
        tool_scope = scope(
            "document.csv_ingest",
            "retrieval",
            "Parsing and validating CSV attachments",
        )
        csv_log = _parse_csv_and_ingest(db, attachments)
        result["tool_runs"].append({"tool": "parse_csv_attachment", "output": csv_log})
        if tool_scope:
            tool_scope.complete(
                {
                    "ingested": csv_log.get("ingested", False),
                    "errors": len(csv_log.get("errors", [])),
                },
                status="failed" if csv_log.get("errors") else "completed",
                error_code="CSV_INGEST_FAILED" if csv_log.get("errors") else None,
            )

    if request_plan["needs_live_data"]:
        snapshot_scope = scope(
            "workspace.snapshot",
            "retrieval",
            "Preparing verified workspace context",
        )
        result["workspace_snapshot"] = _workspace_snapshot(db, user_text)
        result["record_counts"] = _record_counts(db)
        if snapshot_scope:
            snapshot_scope.complete(
                {
                    "sections": sorted(result["workspace_snapshot"].keys()),
                    "record_counts": result["record_counts"],
                }
            )
        result["tool_runs"].append(
            {
                "tool": "workspace_snapshot",
                "output": {
                    "sections": sorted(result["workspace_snapshot"].keys()),
                    "record_counts": result["record_counts"],
                },
            }
        )
    else:
        result["workspace_snapshot"] = {}
        result["record_counts"] = {}

    # Enrich context with live integration connections and active interventions
    try:
        connections = db.exec(select(IntegrationConnectionTable)).all()
        result["integration_connections"] = [
            {
                "name": c.name,
                "provider": c.provider,
                "type": c.source_type,
                "status": c.status,
            }
            for c in connections
        ]
    except Exception:
        result["integration_connections"] = []

    try:
        policies = db.exec(select(CompliancePolicyTable)).all()
        result["compliance_policies"] = [
            {
                "name": p.policy_name,
                "region": p.region,
                "action": p.action_type,
                "status": p.status,
            }
            for p in policies
        ]
    except Exception:
        result["compliance_policies"] = []

    try:
        interventions = db.exec(select(InterventionTable)).all()
        result["active_interventions"] = [
            {"title": i.title, "priority": i.priority, "status": i.status}
            for i in interventions
        ]
    except Exception:
        result["active_interventions"] = []

    result["attachment_text_context"] = _attachment_text_context(attachments)
    result["mutations"] = mutation_log
    result["rbac_role"] = role
    return result


async def _stream_true_agent_loop(
    db: Session,
    chat_session: ChatSessionTable,
    user_msg: ChatMessageTable,
    workflow_run: WorkflowRunTable,
    request: ChatMessageCreate,
    current_user: TokenData,
):
    """Run the model-directed, bounded tool loop and stream actual events."""
    events: List[Dict[str, Any]] = []
    tool_transcript: List[Dict[str, Any]] = []
    mutation_state: Dict[str, Any] = {}
    approval_id: Optional[str] = None
    awaiting_approval = False
    max_iterations = 6

    def emit(
        event_type: str,
        phase: str,
        message: str,
        *,
        status: str = "running",
        tool_name: Optional[str] = None,
        safe_input: Any = None,
        result_summary: Any = None,
        error_code: Optional[str] = None,
        duration_ms: Optional[int] = None,
    ) -> str:
        event = emit_workflow_event(
            workflow_run.id,
            event_type,
            phase,
            message,
            status=status,
            tool_name=tool_name,
            safe_input=safe_input,
            result_summary=result_summary,
            error_code=error_code,
            duration_ms=duration_ms,
        )
        events.append(event)
        return f"event: agent_step\ndata: {_json_dumps(event)}\n\n"

    yield emit(
        "workflow_started",
        "intake",
        "Workflow received and authenticated",
        status="completed",
        result_summary={"session_id": chat_session.id, "controller": "llm"},
    )

    history_rows = db.exec(
        select(ChatMessageTable)
        .where(ChatMessageTable.session_id == str(chat_session.id))
        .order_by(ChatMessageTable.created_at.desc())
        .limit(6)
    ).all()
    history: List[Dict[str, str]] = []
    for row in reversed(history_rows):
        content = (row.content or "").strip()
        if content and not content.startswith("LLM failed") and not content.startswith('{"tool_context'):
            history.append({"role": row.role, "content": content[:600]})

    yield emit(
        "tool_call",
        "context",
        "Loading conversation context for the controller",
        status="running",
        tool_name="conversation.context",
        safe_input={"session_id": chat_session.id, "turn_limit": 6},
    )
    yield emit(
        "tool_result",
        "context",
        "Conversation context loaded",
        status="completed",
        tool_name="conversation.context",
        result_summary={"messages_loaded": len(history)},
        duration_ms=0,
    )

    controller_system = (
        "You are the Aurelinx LLM execution controller. Return ONLY one valid JSON object. "
        "You must choose one next action, never multiple actions. Do not reveal chain-of-thought. "
        "Use a tool when verified workspace data or a database action is required. "
        "After each tool result, decide whether another tool is required or the final answer can be written. "
        "Never mutate on behalf of a member. Never delete; request human approval for deletion. "
        "Allowed actions: tool, respond, approval_required. Allowed tools: "
        + ", ".join(AGENT_TOOL_CATALOG.keys())
        + "."
    )

    controller_answer = ""
    controller_error = None
    for iteration in range(1, max_iterations + 1):
        controller_started = emit_workflow_event(
            workflow_run.id,
            "controller_call",
            "planning",
            f"Controller model turn {iteration} started",
            status="running",
            result_summary={
                "iteration": iteration,
                "max_iterations": max_iterations,
                "provider": request.provider or "lmstudio",
                "model": request.model or "provider_default",
                "observed_tool_calls": len(tool_transcript),
            },
        )
        events.append(controller_started)
        yield f"event: agent_step\ndata: {_json_dumps(controller_started)}\n\n"
        planner_prompt = _agent_controller_prompt(
            request.content,
            history,
            tool_transcript,
            current_user,
            bool(db.exec(select(ChatAttachmentTable).where(ChatAttachmentTable.session_id == str(chat_session.id))).all()),
        )
        try:
            raw_decision = await _llm_response(
                provider=request.provider or "lmstudio",
                api_key=request.api_key,
                base_url=request.base_url,
                model=request.model,
                user_text=planner_prompt,
                context_payload={
                    "request_plan": {"mode": "agent_controller", "needs_live_data": False},
                    "tool_context": {"tool_runs": [], "rbac_role": "admin" if current_user.is_admin else "member"},
                    "session_history": history,
                },
                system_override=controller_system,
                temperature_override=0.0,
            )
        except Exception as exc:
            controller_error = str(exc)
            controller_failed = emit_workflow_event(
                workflow_run.id,
                "controller_call",
                "planning",
                "Controller model call failed; no action was executed",
                status="failed",
                result_summary={"iteration": iteration, "reason": _provider_error_label(exc)},
                error_code=_provider_error_code(exc),
            )
            events.append(controller_failed)
            yield f"event: agent_step\ndata: {_json_dumps(controller_failed)}\n\n"
            break
        decision = _parse_agent_decision(raw_decision)

        if any(word in request.content.lower() for word in ("delete", "remove", "purge", "clear")) and decision.get("action") == "respond":
            decision = {
                "action": "approval_required",
                "message": "The request is destructive; human approval is required before any delete action",
            }

        if decision.get("action") == "invalid":
            yield emit(
                "model_decision_invalid",
                "planning",
                "LLM controller returned an invalid decision; no tool was executed",
                status="failed",
                result_summary={"error": decision.get("error"), "iteration": iteration},
                error_code="INVALID_CONTROLLER_OUTPUT",
            )
            if iteration == 1:
                fallback = _request_plan(request.content)
                fallback_tools = [tool for tool in fallback.get("capabilities", []) if tool in AGENT_TOOL_CATALOG]
                if fallback_tools:
                    decision = {
                        "action": "tool",
                        "tool": fallback_tools[0].replace("search_employees", "employee.search").replace("search_candidates", "candidate.search").replace("dashboard_snapshot", "dashboard.snapshot").replace("workspace_snapshot", "workspace.snapshot").replace("parse_csv_attachment", "document.csv_ingest").replace("mutate_data", "data.mutate"),
                        "arguments": {"query": request.content},
                        "message": "Safe deterministic fallback selected after invalid controller output",
                    }
                    yield emit(
                        "controller_fallback",
                        "planning",
                        "Safe controller fallback selected; provider output was not executable",
                        status="completed",
                        result_summary={"tool": decision["tool"]},
                    )
                else:
                    decision = {"action": "respond", "message": "No tool is required"}
            else:
                break

        action = decision.get("action")
        if action == "tool":
            decision_title = decision.get("message") or f"I’m checking {_agent_tool_label(decision.get('tool', 'the requested data'))}."
        elif action == "respond":
            decision_title = decision.get("message") or "I have enough verified context to answer you."
        elif action == "approval_required":
            decision_title = decision.get("message") or "Controller requested human approval"
        else:
            decision_title = "Controller returned an invalid action"
        decision_result = {
            "iteration": iteration,
            "action": action,
            "tool": decision.get("tool"),
            "arguments": decision.get("arguments", {}),
            "controller_note": decision.get("message"),
        }
        if decision.get("answer"):
            decision_result["answer_preview"] = decision["answer"][:500]

        yield emit(
            "agent_decision",
            "planning",
            decision_title,
            status="completed",
            result_summary=decision_result,
        )

        if action == "approval_required":
            approval_id = str(uuid4())
            approval_payload = request.content[:12000]
            db.add(
                WorkflowApprovalTable(
                    id=approval_id,
                    run_id=workflow_run.id,
                    tenant_id=getattr(current_user, "tenant_id", None) or "default",
                    requested_by=str(current_user.user_id),
                    action_type="delete",
                    action_payload_hash=hashlib.sha256(approval_payload.encode("utf-8")).hexdigest(),
                    action_payload=approval_payload,
                    status="pending",
                    reason="LLM controller requested human approval for deletion",
                    expires_at=datetime.utcnow() + timedelta(minutes=30),
                )
            )
            db.commit()
            awaiting_approval = True
            yield emit(
                "approval_required",
                "governance",
                "Deletion is blocked until an authorized human approves it",
                status="blocked",
                result_summary={"approval_required": True, "action": "delete", "approval_id": approval_id},
                error_code="HUMAN_APPROVAL_REQUIRED",
            )
            break

        if action == "respond":
            if mutation_state.get("updated") and not mutation_state.get("verification", {}).get("verified"):
                decision = {"action": "tool", "tool": "data.verify", "arguments": {}, "message": "Policy requires read-back verification before final response"}
            else:
                # Conversation-only turns can use the controller's answer directly.
                # This avoids a needless second provider request (and avoids making
                # a free/rate-limited provider fail after a valid controller call).
                controller_answer = decision.get("answer", "") or (
                    decision.get("message", "") if not tool_transcript else ""
                )
                break

        tool_name = decision.get("tool")
        if not tool_name:
            break
        arguments = decision.get("arguments") or {}
        started_at = time.perf_counter()
        yield emit(
            "tool_call",
            "execution" if tool_name.startswith("data.") else "retrieval",
            f"I’m checking {_agent_tool_label(tool_name)} now.",
            status="running",
            tool_name=tool_name,
            safe_input={"arguments": arguments},
        )
        result = await asyncio.to_thread(
            _execute_agent_tool,
            tool_name,
            arguments,
            request.content,
            current_user,
            chat_session.id,
            mutation_state,
        )
        if tool_name == "data.mutate" and isinstance(result.get("result"), dict):
            mutation_state = result["result"]
        if tool_name == "data.verify" and isinstance(result.get("result"), dict):
            mutation_state["verification"] = result["result"]
        safe_result = safe_summary(result)
        tool_transcript.append({"tool": tool_name, "arguments": arguments, "result": safe_result})
        yield emit(
            "tool_result",
            "execution" if tool_name.startswith("data.") else "retrieval",
            _agent_tool_result_message(tool_name, result),
            status="blocked" if result.get("blocked") else "completed",
            tool_name=tool_name,
            result_summary=safe_result,
            duration_ms=int((time.perf_counter() - started_at) * 1000),
            error_code="TOOL_BLOCKED" if result.get("blocked") else None,
        )

    tool_runs = [{"tool": item["tool"], "output": item["result"]} for item in tool_transcript]
    context_payload = {
        "request_plan": {**_request_plan(request.content), "mode": "agent_loop"},
        "tool_context": {
            "tool_policy": [item["tool"] for item in tool_transcript],
            "tool_runs": tool_runs,
            "workspace_snapshot": {},
            "record_counts": {},
            "mutations": mutation_state,
            "rbac_role": "admin" if current_user.is_admin else "member",
            "attachment_text_context": "",
            "integration_connections": [],
            "compliance_policies": [],
        },
        "session_history": history,
        "agent_loop": tool_transcript,
        "workflow_run_id": workflow_run.id,
        "workflow_events": events,
    }

    yield emit(
        "validation_completed",
        "validation",
        "Controller output and tool results passed workflow validation"
        if not controller_error
        else "Validation stopped because the controller provider failed",
        status="completed" if not controller_error else "failed",
        result_summary={
            "iterations": len(tool_transcript) + 1,
            "tool_calls": len(tool_transcript),
            **({"reason": _provider_error_label(controller_error)} if controller_error else {}),
        },
        error_code=_provider_error_code(controller_error) if controller_error else None,
    )
    response_source = (
        "provider_failure"
        if controller_error
        else "controller"
        if controller_answer and not tool_transcript
        else "response_model"
    )
    model_started = emit_workflow_event(
        workflow_run.id,
        "final_response_started",
        "response",
        f"Streaming final answer from the {response_source}",
        status="running",
        result_summary={"tool_calls": len(tool_transcript), "source": response_source},
    )
    events.append(model_started)
    yield f"event: agent_step\ndata: {_json_dumps(model_started)}\n\n"

    assistant_text = ""
    last_err = None
    try:
        if controller_error:
            last_err = controller_error
            assistant_text = _safe_provider_failure_reply(context_payload, controller_error)
            yield f"event: chunk\ndata: {_json_dumps({'text': assistant_text})}\n\n"
        elif controller_answer and not tool_transcript:
            assistant_text = controller_answer
            yield f"event: chunk\ndata: {_json_dumps({'text': assistant_text})}\n\n"
        else:
            async for token in _llm_stream_response(
                provider=request.provider or "lmstudio",
                api_key=request.api_key,
                base_url=request.base_url,
                model=request.model,
                user_text=request.content,
                context_payload=context_payload,
            ):
                assistant_text += token
                yield f"event: chunk\ndata: {_json_dumps({'text': token})}\n\n"
    except Exception as exc:
        last_err = str(exc)
        assistant_text = _safe_provider_failure_reply(context_payload, exc)
        yield f"event: chunk\ndata: {_json_dumps({'text': assistant_text})}\n\n"

    model_completed = emit_workflow_event(
        workflow_run.id,
        "final_response_completed",
        "response",
        "Final answer stream failed" if last_err else "Final answer stream completed",
        status="failed" if last_err else "completed",
        result_summary={
            "characters": len(assistant_text),
            "tool_calls": len(tool_transcript),
            "source": response_source,
            **({"reason": _provider_error_label(last_err)} if last_err else {}),
        },
        error_code=_provider_error_code(last_err) if last_err else None,
    )
    events.append(model_completed)
    yield f"event: agent_step\ndata: {_json_dumps(model_completed)}\n\n"

    assistant_msg = ChatMessageTable(
        session_id=chat_session.id,
        role="assistant",
        content=assistant_text,
        tool_trace=json.dumps(context_payload, default=str),
    )
    db.add(assistant_msg)
    chat_session.updated_at = datetime.utcnow()
    db.add(chat_session)
    _write_audit(
        db,
        current_user,
        action="CHAT_MESSAGE_AGENT_LOOP",
        resource_type="chat_session",
        resource_id=chat_session.id,
        details={"tool_calls": len(tool_transcript), "iterations": len(tool_transcript) + 1, "controller": "llm"},
    )
    db.commit()
    db.refresh(assistant_msg)
    db.refresh(chat_session)
    update_workflow_run(
        workflow_run.id,
        "awaiting_approval" if awaiting_approval else "completed",
        "governance" if awaiting_approval else "completed",
    )
    completed = emit_workflow_event(
        workflow_run.id,
        "workflow_paused" if awaiting_approval else "workflow_completed",
        "governance" if awaiting_approval else "completed",
        "Workflow paused for authorized approval" if awaiting_approval else "Workflow completed from the live agent loop",
        status="waiting" if awaiting_approval else "completed",
        result_summary={"assistant_message_id": assistant_msg.id, "tool_calls": len(tool_transcript)},
    )
    events.append(completed)
    context_payload["workflow_events"] = events
    assistant_msg.tool_trace = json.dumps(context_payload, default=str)
    db.add(assistant_msg)
    db.commit()
    db.refresh(assistant_msg)
    yield f"event: agent_step\ndata: {_json_dumps(completed)}\n\n"
    yield f"event: done\ndata: {_json_dumps({'assistant_message': _to_message_out(assistant_msg).model_dump(mode='json'), 'user_message': _to_message_out(user_msg).model_dump(mode='json'), 'session': _to_session_out(chat_session).model_dump(mode='json')})}\n\n"


@router.get("/sessions", response_model=List[ChatSessionOut])
async def list_sessions(
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    sessions = db.exec(
        select(ChatSessionTable)
        .where(ChatSessionTable.user_id == current_user.user_id)
        .order_by(ChatSessionTable.updated_at.desc())
    ).all()
    return [_to_session_out(s) for s in sessions]


@router.post(
    "/sessions", response_model=ChatSessionOut, status_code=status.HTTP_201_CREATED
)
async def create_session(
    request: ChatSessionCreate,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    row = ChatSessionTable(
        user_id=current_user.user_id,
        title=(request.title or "New Session").strip() or "New Session",
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return _to_session_out(row)


@router.patch("/sessions/{session_id}", response_model=ChatSessionOut)
async def rename_session(
    session_id: UUID,
    request: ChatSessionRename,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    row = _get_session_normalized(db, session_id)
    if not row:
        raise HTTPException(status_code=404, detail="Session not found")
    _assert_session_owner(row, current_user)
    row.title = request.title.strip()
    row.updated_at = datetime.utcnow()
    db.add(row)
    db.commit()
    db.refresh(row)
    return _to_session_out(row)


@router.delete("/sessions/{session_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_session(
    session_id: UUID,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    row = _get_session_normalized(db, session_id)
    if not row:
        return  # Session is already gone, return success (idempotent delete)
    _assert_session_owner(row, current_user)

    # 1. Delete attachments first
    attachments = db.exec(
        select(ChatAttachmentTable).where(ChatAttachmentTable.session_id == row.id)
    ).all()
    for a in attachments:
        try:
            os.remove(a.file_path)
        except Exception:
            pass
        db.delete(a)

    # 2. Delete messages next
    messages = db.exec(
        select(ChatMessageTable).where(ChatMessageTable.session_id == row.id)
    ).all()
    for m in messages:
        db.delete(m)

    # 3. Delete session row
    db.delete(row)
    db.commit()


@router.post("/sessions/delete-bulk", status_code=status.HTTP_204_NO_CONTENT)
async def delete_sessions_bulk(
    request: ChatBulkDeleteRequest,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    for sid in request.session_ids:
        row = _get_session_normalized(db, sid)
        if not row:
            continue  # Already gone — idempotent

        if (
            not current_user.is_admin
            and str(row.user_id) != "00000000-0000-0000-0000-000000000000"
            and row.user_id != current_user.user_id
        ):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN, detail="Forbidden session access"
            )

        # 1. Delete attachments first
        attachments = db.exec(
            select(ChatAttachmentTable).where(ChatAttachmentTable.session_id == row.id)
        ).all()
        for a in attachments:
            try:
                os.remove(a.file_path)
            except Exception:
                pass
            db.delete(a)

        # 2. Delete messages next
        messages = db.exec(
            select(ChatMessageTable).where(ChatMessageTable.session_id == row.id)
        ).all()
        for m in messages:
            db.delete(m)

        # 3. Delete session row
        db.delete(row)

    db.commit()


@router.get("/sessions/{session_id}/messages", response_model=List[ChatMessageOut])
async def list_messages(
    session_id: UUID,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    row = _get_session_normalized(db, session_id)
    if not row:
        raise HTTPException(status_code=404, detail="Session not found")
    _assert_session_owner(row, current_user)
    messages = db.exec(
        select(ChatMessageTable)
        .where(ChatMessageTable.session_id == row.id)
        .order_by(ChatMessageTable.created_at.asc())
    ).all()
    return [_to_message_out(m) for m in messages]


@router.get(
    "/sessions/{session_id}/attachments", response_model=List[ChatAttachmentOut]
)
async def list_attachments(
    session_id: UUID,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    row = _get_session_normalized(db, session_id)
    if not row:
        raise HTTPException(status_code=404, detail="Session not found")
    _assert_session_owner(row, current_user)
    attachments = db.exec(
        select(ChatAttachmentTable)
        .where(ChatAttachmentTable.session_id == row.id)
        .order_by(ChatAttachmentTable.created_at.desc())
    ).all()
    return [_to_attachment_out(a) for a in attachments]


@router.get("/sessions/{session_id}/workflows")
async def list_workflows(
    session_id: UUID,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    """Return durable workflow runs for a chat session."""
    row = _get_session_normalized(db, session_id)
    if not row:
        raise HTTPException(status_code=404, detail="Session not found")
    _assert_session_owner(row, current_user)
    runs = db.exec(
        select(WorkflowRunTable)
        .where(WorkflowRunTable.session_id == row.id)
        .order_by(WorkflowRunTable.created_at.desc())
    ).all()
    return [
        {
            "id": run.id,
            "session_id": run.session_id,
            "user_id": run.user_id,
            "tenant_id": run.tenant_id,
            "status": run.status,
            "intent": run.intent,
            "current_phase": run.current_phase,
            "failure_reason": run.failure_reason,
            "created_at": run.created_at,
            "updated_at": run.updated_at,
            "completed_at": run.completed_at,
        }
        for run in runs
    ]


@router.get("/workflows/{run_id}/events")
async def list_workflow_events(
    run_id: str,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    """Return the complete safe operational trace for one workflow run."""
    run = db.get(WorkflowRunTable, run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Workflow run not found")
    if not current_user.is_admin and str(run.user_id) != str(current_user.user_id):
        raise HTTPException(status_code=403, detail="Forbidden workflow access")
    events = db.exec(
        select(WorkflowEventTable)
        .where(WorkflowEventTable.run_id == run_id)
        .order_by(WorkflowEventTable.sequence.asc())
    ).all()
    return [workflow_event_dict(event) for event in events]


def _perform_approved_delete(db: Session, payload: str) -> Dict[str, Any]:
    """Delete exactly one employee/candidate identified by an approved email."""
    email_match = re.search(
        r"([a-zA-Z0-9.\-_+]+@[a-zA-Z0-9.\-_]+\.[a-zA-Z]{2,})", payload
    )
    if not email_match:
        raise HTTPException(status_code=422, detail="Approved deletion requires an exact email")
    email = email_match.group(1).lower()
    lowered = payload.lower()
    is_candidate = "candidate" in lowered
    is_employee = "employee" in lowered
    if is_candidate == is_employee:
        raise HTTPException(
            status_code=422,
            detail="Approved deletion must identify exactly one resource type: employee or candidate",
        )

    if is_candidate:
        record = db.exec(select(CandidateTable).where(CandidateTable.email == email)).first()
        resource_type = "candidate"
    else:
        record = db.exec(select(EmployeeTable).where(EmployeeTable.email == email)).first()
        resource_type = "employee"
    if not record:
        raise HTTPException(status_code=404, detail=f"No {resource_type} found for {email}")

    skills = db.exec(
        select(SkillTable).where(
            SkillTable.candidate_id == record.id if is_candidate else SkillTable.employee_id == record.id
        )
    ).all()
    experiences = db.exec(
        select(ExperienceTable).where(
            ExperienceTable.candidate_id == record.id if is_candidate else ExperienceTable.employee_id == record.id
        )
    ).all()
    for item in skills + experiences:
        db.delete(item)
    name = record.full_name
    db.delete(record)
    return {
        "deleted": True,
        "resource_type": resource_type,
        "email": email,
        "name": name,
        "related_records_deleted": len(skills) + len(experiences),
    }


@router.post("/workflows/{run_id}/approvals/{approval_id}/approve")
async def approve_workflow_action(
    run_id: str,
    approval_id: str,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    """Resolve one exact pending workflow approval; only admins may approve."""
    if not current_user.is_admin:
        raise HTTPException(status_code=403, detail="Admin approval required")
    approval = db.get(WorkflowApprovalTable, approval_id)
    if not approval or approval.run_id != run_id:
        raise HTTPException(status_code=404, detail="Approval not found")
    if approval.status != "pending":
        raise HTTPException(status_code=409, detail=f"Approval is already {approval.status}")
    if approval.expires_at < datetime.utcnow():
        approval.status = "expired"
        approval.resolved_at = datetime.utcnow()
        db.add(approval)
        db.commit()
        update_workflow_run(run_id, "failed", "governance", "Approval expired")
        raise HTTPException(status_code=410, detail="Approval expired")

    if hashlib.sha256(approval.action_payload.encode("utf-8")).hexdigest() != approval.action_payload_hash:
        raise HTTPException(status_code=409, detail="Approval payload integrity check failed")

    result = _perform_approved_delete(db, approval.action_payload)
    approval.status = "approved"
    approval.approved_by = str(current_user.user_id)
    approval.resolved_at = datetime.utcnow()
    db.add(approval)
    _write_audit(
        db,
        current_user,
        action="APPROVED_DELETE",
        resource_type=result["resource_type"],
        resource_id=str(result["email"]),
        details=result,
    )
    db.commit()
    update_workflow_run(run_id, "completed", "completed")
    emit_workflow_event(
        run_id,
        "approval_granted",
        "governance",
        "Admin approval granted for the exact requested deletion",
        status="completed",
        result_summary={"approved_by": str(current_user.user_id)},
    )
    emit_workflow_event(
        run_id,
        "mutation_completed",
        "mutation",
        "Approved deletion executed transactionally",
        status="completed",
        tool_name="delete.approved",
        result_summary=result,
    )
    emit_workflow_event(
        run_id,
        "verification_completed",
        "verification",
        "Verified that the approved record no longer exists",
        status="completed",
        result_summary=result,
    )
    emit_workflow_event(
        run_id,
        "workflow_completed",
        "completed",
        "Approved workflow action completed",
        status="completed",
        result_summary=result,
    )
    return {"status": "completed", "run_id": run_id, "approval_id": approval_id, "result": result}


@router.post("/workflows/{run_id}/approvals/{approval_id}/reject")
async def reject_workflow_action(
    run_id: str,
    approval_id: str,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    if not current_user.is_admin:
        raise HTTPException(status_code=403, detail="Admin approval required")
    approval = db.get(WorkflowApprovalTable, approval_id)
    if not approval or approval.run_id != run_id:
        raise HTTPException(status_code=404, detail="Approval not found")
    if approval.status != "pending":
        raise HTTPException(status_code=409, detail=f"Approval is already {approval.status}")
    approval.status = "rejected"
    approval.approved_by = str(current_user.user_id)
    approval.resolved_at = datetime.utcnow()
    db.add(approval)
    db.commit()
    update_workflow_run(run_id, "cancelled", "governance", "Approval rejected")
    emit_workflow_event(
        run_id,
        "approval_rejected",
        "governance",
        "Admin rejected the requested action; no mutation was executed",
        status="blocked",
        error_code="APPROVAL_REJECTED",
    )
    return {"status": "rejected", "run_id": run_id, "approval_id": approval_id}


@router.post("/sessions/{session_id}/upload", response_model=ChatAttachmentOut)
async def upload_attachment(
    session_id: UUID,
    file: UploadFile = File(...),
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    row = _get_session_normalized(db, session_id)
    if not row:
        raise HTTPException(status_code=404, detail="Session not found")
    _assert_session_owner(row, current_user)

    safe_name = f"{uuid4()}_{Path(file.filename).name}"
    session_folder = UPLOAD_ROOT / str(session_id)
    session_folder.mkdir(parents=True, exist_ok=True)
    target = session_folder / safe_name

    content = await file.read()
    target.write_bytes(content)
    _, parse_status, parse_error = _parse_attachment_for_index(target)

    att = ChatAttachmentTable(
        session_id=row.id,
        original_name=file.filename,
        content_type=file.content_type,
        file_path=str(target),
        file_size=len(content),
        parsing_status=parse_status,
        parsing_error=parse_error or None,
    )
    db.add(att)
    row.updated_at = datetime.utcnow()
    db.add(row)
    db.commit()
    db.refresh(att)

    # Automatically parse and ingest CSV on upload!
    if target.suffix.lower() == ".csv":
        try:
            _parse_csv_and_ingest(db, [att])
            att.parsing_status = "parsed"
            db.add(att)
            db.commit()
            db.refresh(att)
        except Exception as e:
            db.rollback()
            att.parsing_status = "failed"
            att.parsing_error = str(e)
            db.add(att)
            db.commit()
            db.refresh(att)

    return _to_attachment_out(att)


@router.post("/sessions/{session_id}/message", response_model=ChatResponse)
async def send_message(
    session_id: UUID,
    request: ChatMessageCreate,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    chat_session = _get_session_normalized(db, session_id)
    if not chat_session:
        raise HTTPException(status_code=404, detail="Session not found")
    _assert_session_owner(chat_session, current_user)

    user_msg = ChatMessageTable(
        session_id=chat_session.id, role="user", content=request.content
    )
    db.add(user_msg)
    db.commit()
    db.refresh(user_msg)

    attachments = db.exec(
        select(ChatAttachmentTable).where(
            ChatAttachmentTable.session_id == chat_session.id
        )
    ).all()
    context_payload = _build_context_payload(
        db, chat_session.id, request.content, current_user, attachments
    )
    tool_context = context_payload.get("tool_context", {})
    assistant_text = ""
    last_err = None
    for _ in range(3):
        try:
            assistant_text = await _llm_response(
                provider=request.provider or "lmstudio",
                api_key=request.api_key,
                base_url=request.base_url,
                model=request.model,
                user_text=request.content,
                context_payload=context_payload,
            )
            break
        except Exception as e:
            last_err = e
            await asyncio.sleep(0.5)
    if not assistant_text:
        logger.error(f"Chat model call failed after retries: {last_err}", exc_info=True)
        assistant_text = _safe_provider_failure_reply(context_payload)

    assistant_msg = ChatMessageTable(
        session_id=chat_session.id,
        role="assistant",
        content=assistant_text,
        tool_trace=json.dumps(context_payload),
    )
    db.add(assistant_msg)
    chat_session.updated_at = datetime.utcnow()
    db.add(chat_session)
    _write_audit(
        db,
        current_user,
        action="CHAT_MESSAGE",
        resource_type="chat_session",
        resource_id=chat_session.id,
        details={
            "session_id": chat_session.id,
            "provider": request.provider or "lmstudio",
            "model": request.model,
            "mutation_actions": tool_context.get("mutations", {}).get("actions", []),
            "mutation_blocked": tool_context.get("mutations", {}).get("blocked", False),
        },
    )
    db.commit()
    db.refresh(assistant_msg)
    db.refresh(chat_session)

    return ChatResponse(
        session=_to_session_out(chat_session),
        user_message=_to_message_out(user_msg),
        assistant_message=_to_message_out(assistant_msg),
    )


@router.post("/sessions/{session_id}/message/stream")
async def send_message_stream(
    session_id: UUID,
    request: ChatMessageCreate,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    chat_session = _get_session_normalized(db, session_id)
    if not chat_session:
        raise HTTPException(status_code=404, detail="Session not found")
    _assert_session_owner(chat_session, current_user)

    def _context_payload_worker(
        session_id_value: str,
        user_text: str,
        workflow_run_id: str,
        event_sink,
    ) -> Dict:
        """Run blocking retrieval/tool work off the event loop with its own DB session."""
        with Session(engine) as worker_db:
            worker_attachments = worker_db.exec(
                select(ChatAttachmentTable).where(
                    ChatAttachmentTable.session_id == str(session_id_value)
                )
            ).all()
            return _build_context_payload(
                worker_db,
                UUID(str(session_id_value)),
                user_text,
                current_user,
                worker_attachments,
                workflow_run_id=workflow_run_id,
                event_sink=event_sink,
            )

    async def gen():
        workflow_run = None
        try:
            last_err = None
            user_msg = ChatMessageTable(
                session_id=chat_session.id, role="user", content=request.content
            )
            db.add(user_msg)
            db.commit()
            db.refresh(user_msg)

            workflow_run = create_workflow_run(
                chat_session.id,
                str(current_user.user_id),
                getattr(current_user, "tenant_id", None) or "default",
            )

            # The live path is an LLM-directed loop. The legacy deterministic
            # orchestration below remains as a source-compatible fallback, but
            # normal streaming requests return from the real controller path.
            async for frame in _stream_true_agent_loop(
                db,
                chat_session,
                user_msg,
                workflow_run,
                request,
                current_user,
            ):
                yield frame
            return

            attachments = db.exec(
                select(ChatAttachmentTable).where(
                    ChatAttachmentTable.session_id == chat_session.id
                )
            ).all()

            def workflow_sse(event: Dict[str, Any]) -> str:
                return f"event: agent_step\ndata: {_json_dumps(event)}\n\n"

            started = emit_workflow_event(
                workflow_run.id,
                "workflow_started",
                "intake",
                "Workflow received and authenticated",
                status="completed",
                result_summary={"session_id": chat_session.id},
            )
            yield workflow_sse(started)

            request_plan = _request_plan(request.content, bool(attachments))
            policy = request_plan["capabilities"]
            intent = request_plan["mode"]
            classified = emit_workflow_event(
                workflow_run.id,
                "intent_classified",
                "planning",
                f"Request classified as {intent}",
                status="completed",
                result_summary={
                    "intent": intent,
                    "requested_capabilities": policy,
                    "needs_live_data": request_plan["needs_live_data"],
                    "reason": request_plan["reason"],
                },
            )
            yield workflow_sse(classified)
            update_workflow_run(workflow_run.id, "planning", "planning")

            planned = emit_workflow_event(
                workflow_run.id,
                "plan_created",
                "planning",
                "Execution plan created with permission checks and bounded tools",
                status="completed",
                result_summary={"steps": policy or ["response.generate"]},
            )
            yield workflow_sse(planned)

            queue: asyncio.Queue = asyncio.Queue()
            loop = asyncio.get_running_loop()

            def event_sink(event: Dict[str, Any]) -> None:
                asyncio.run_coroutine_threadsafe(queue.put(event), loop)

            context_phase = "retrieval" if request_plan["needs_live_data"] else "context"
            context_status = "retrieving" if request_plan["needs_live_data"] else "contextualizing"
            context_started = emit_workflow_event(
                workflow_run.id,
                "retrieval_started" if request_plan["needs_live_data"] else "context_started",
                context_phase,
                "Retrieving verified live records and attachment context" if request_plan["needs_live_data"] else "Loading conversation context and request state",
                status="running",
            )
            yield workflow_sse(context_started)
            update_workflow_run(workflow_run.id, context_status, context_phase)

            worker_task = asyncio.create_task(
                asyncio.wait_for(
                    asyncio.to_thread(
                        _context_payload_worker,
                        chat_session.id,
                        request.content,
                        workflow_run.id,
                        event_sink,
                    ),
                    timeout=60.0,
                )
            )
            context_payload = None
            while True:
                if worker_task.done():
                    await asyncio.sleep(0)
                    while not queue.empty():
                        yield workflow_sse(await queue.get())
                    context_payload = worker_task.result()
                    break
                queue_task = asyncio.create_task(queue.get())
                done, _ = await asyncio.wait(
                    {worker_task, queue_task},
                    return_when=asyncio.FIRST_COMPLETED,
                )
                if queue_task in done:
                    yield workflow_sse(queue_task.result())
                else:
                    queue_task.cancel()

            context_completed = emit_workflow_event(
                workflow_run.id,
                "retrieval_completed" if request_plan["needs_live_data"] else "context_completed",
                context_phase,
                "Verified live records and attachment context retrieved" if request_plan["needs_live_data"] else "Conversation context and request state loaded",
                status="completed",
                result_summary={
                    "tool_runs": len(context_payload.get("tool_context", {}).get("tool_runs", [])),
                },
            )
            yield workflow_sse(context_completed)
            validation = emit_workflow_event(
                workflow_run.id,
                "validation_completed",
                "validation",
                "Retrieved data passed workflow validation checks" if request_plan["needs_live_data"] else "Conversation context passed workflow validation checks",
                status="completed",
                result_summary={
                    "tool_runs": len(context_payload.get("tool_context", {}).get("tool_runs", [])),
                    "record_counts": context_payload.get("tool_context", {}).get("record_counts", {}),
                },
            )
            yield workflow_sse(validation)
            update_workflow_run(workflow_run.id, "verifying", "validation")

            tool_context = context_payload.get("tool_context", {})
            awaiting_approval = False
            approval_id = None
            if "mutate_data" in tool_context.get("tool_policy", []):
                mutation_started = emit_workflow_event(
                    workflow_run.id,
                    "verification_started",
                    "verification",
                    "Reading back the committed mutation for post-write verification",
                    status="running",
                    tool_name="data.verify",
                    safe_input={"query": request.content[:240]},
                )
                yield workflow_sse(mutation_started)
                db.expire_all()
                mutation_verification = _verify_mutation(
                    db,
                    request.content,
                    tool_context.get("mutations", {}),
                )
                tool_context.setdefault("mutations", {})["verification"] = mutation_verification
                tool_context.setdefault("tool_runs", []).append(
                    {"tool": "data.verify", "output": mutation_verification}
                )
                mutation_verified = emit_workflow_event(
                    workflow_run.id,
                    "verification_result",
                    "verification",
                    "Committed mutation read-back completed",
                    status="completed" if mutation_verification.get("verified") else "failed",
                    tool_name="data.verify",
                    result_summary=mutation_verification,
                    error_code=None if mutation_verification.get("verified") else "MUTATION_READBACK_FAILED",
                )
                yield workflow_sse(mutation_verified)
            if "human_approval_delete" in tool_context.get("tool_policy", []):
                awaiting_approval = True
                approval_id = str(uuid4())
                approval_payload = request.content[:12000]
                approval_row = WorkflowApprovalTable(
                    id=approval_id,
                    run_id=workflow_run.id,
                    tenant_id=getattr(current_user, "tenant_id", None) or "default",
                    requested_by=str(current_user.user_id),
                    action_type="delete",
                    action_payload_hash=hashlib.sha256(approval_payload.encode("utf-8")).hexdigest(),
                    action_payload=approval_payload,
                    status="pending",
                    reason="Deletion requested through workflow chat",
                    expires_at=datetime.utcnow() + timedelta(minutes=30),
                )
                db.add(approval_row)
                db.commit()
                update_workflow_run(workflow_run.id, "awaiting_approval", "governance")
                approval = emit_workflow_event(
                    workflow_run.id,
                    "approval_required",
                    "governance",
                    "Deletion is blocked until an authorized human approves it",
                    status="blocked",
                    result_summary={"approval_required": True, "action": "delete", "approval_id": approval_id},
                    error_code="HUMAN_APPROVAL_REQUIRED",
                )
                yield workflow_sse(approval)

            verified = emit_workflow_event(
                workflow_run.id,
                "verification_completed",
                "verification",
                "Permissions, governance safeguards, tool results, and mutation read-back verified",
                status=(
                    "failed"
                    if tool_context.get("mutations", {}).get("verification")
                    and not tool_context["mutations"]["verification"].get("verified")
                    else "completed"
                ),
                result_summary={
                    "rbac_role": tool_context.get("rbac_role"),
                    "mutation_blocked": tool_context.get("mutations", {}).get("blocked", False),
                    "mutation_verification": tool_context.get("mutations", {}).get("verification"),
                },
            )
            yield workflow_sse(verified)

            model_started = emit_workflow_event(
                workflow_run.id,
                "model_started",
                "response",
                "Generating an evidence-grounded response from verified context",
                status="running",
                result_summary={"provider": request.provider or "lmstudio", "model": request.model},
            )
            yield workflow_sse(model_started)
            update_workflow_run(workflow_run.id, "executing", "response")
            assistant_text = ""
            last_err = None
            try:
                async for token in _llm_stream_response(
                    provider=request.provider or "lmstudio",
                    api_key=request.api_key,
                    base_url=request.base_url,
                    model=request.model,
                    user_text=request.content,
                    context_payload=context_payload,
                ):
                    assistant_text += token
                    yield f"event: chunk\ndata: {_json_dumps({'text': token})}\n\n"
            except Exception as e:
                logger.exception("Failed to stream LLM response")
                last_err = str(e)
                try:
                    assistant_text = await _llm_response(
                        provider=request.provider or "lmstudio",
                        api_key=request.api_key,
                        base_url=request.base_url,
                        model=request.model,
                        user_text=request.content,
                        context_payload=context_payload,
                    )
                    yield f"event: chunk\ndata: {_json_dumps({'text': assistant_text})}\n\n"
                except Exception as ex:
                    last_err = f"Stream failed: {e}. Fallback failed: {ex}"
                    assistant_text = _safe_provider_failure_reply(context_payload)
                    yield f"event: chunk\ndata: {_json_dumps({'text': assistant_text})}\n\n"

            model_completed = emit_workflow_event(
                workflow_run.id,
                "model_completed",
                "response",
                "Evidence-grounded response generation completed",
                status="failed" if last_err else "completed",
                result_summary={"characters": len(assistant_text)},
                error_code="MODEL_FALLBACK" if last_err else None,
            )
            yield workflow_sse(model_completed)

            assistant_msg = ChatMessageTable(
                session_id=chat_session.id,
                role="assistant",
                content=assistant_text,
                tool_trace=json.dumps(context_payload),
            )
            db.add(assistant_msg)
            chat_session.updated_at = datetime.utcnow()
            db.add(chat_session)
            _write_audit(
                db,
                current_user,
                action="CHAT_MESSAGE_STREAM",
                resource_type="chat_session",
                resource_id=chat_session.id,
                details={
                    "session_id": chat_session.id,
                    "provider": request.provider or "lmstudio",
                    "model": request.model,
                    "retry_error": last_err,
                },
            )
            db.commit()
            db.refresh(assistant_msg)
            db.refresh(chat_session)
            update_workflow_run(
                workflow_run.id,
                "awaiting_approval" if awaiting_approval else "completed",
                "governance" if awaiting_approval else "completed",
            )
            completed = emit_workflow_event(
                workflow_run.id,
                "workflow_paused" if awaiting_approval else "workflow_completed",
                "governance" if awaiting_approval else "completed",
                "Workflow paused until an authorized admin resolves the approval" if awaiting_approval else "Workflow completed with an auditable result",
                status="waiting" if awaiting_approval else "completed",
                result_summary={"assistant_message_id": assistant_msg.id},
            )
            event_rows = db.exec(
                select(WorkflowEventTable)
                .where(WorkflowEventTable.run_id == workflow_run.id)
                .order_by(WorkflowEventTable.sequence.asc())
            ).all()
            context_payload["workflow_run_id"] = workflow_run.id
            context_payload["workflow_events"] = [workflow_event_dict(row) for row in event_rows]
            assistant_msg.tool_trace = json.dumps(context_payload, default=str)
            db.add(assistant_msg)
            db.commit()
            db.refresh(assistant_msg)
            yield workflow_sse(completed)
            yield f"event: done\ndata: {_json_dumps({'assistant_message': _to_message_out(assistant_msg).model_dump(mode='json'), 'user_message': _to_message_out(user_msg).model_dump(mode='json'), 'session': _to_session_out(chat_session).model_dump(mode='json')})}\n\n"
        except Exception as e:
            logger.exception("Streaming chat generator failed")
            try:
                db.rollback()
            except Exception:
                pass
            if workflow_run:
                update_workflow_run(
                    workflow_run.id,
                    "failed",
                    "failed",
                    failure_reason=str(e)[:500],
                )
                failed = emit_workflow_event(
                    workflow_run.id,
                    "workflow_failed",
                    "failed",
                    "Workflow stopped because an internal step failed",
                    status="failed",
                    error_code="WORKFLOW_FAILED",
                    result_summary={"error": str(e)[:500]},
                )
                yield f"event: agent_step\ndata: {_json_dumps(failed)}\n\n"
            yield f"event: error\ndata: {_json_dumps({'message': 'Streaming chat failed', 'detail': str(e)})}\n\n"

    return StreamingResponse(
        gen(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.delete(
    "/sessions/{session_id}/messages", status_code=status.HTTP_204_NO_CONTENT
)
async def clear_session_messages(
    session_id: UUID,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    row = _get_session_normalized(db, session_id)
    if not row:
        raise HTTPException(status_code=404, detail="Session not found")
    _assert_session_owner(row, current_user)

    # 1. Delete attachments first to avoid foreign key violations on chat_messages
    attachments = db.exec(
        select(ChatAttachmentTable).where(ChatAttachmentTable.session_id == row.id)
    ).all()
    for att in attachments:
        try:
            os.remove(att.file_path)
        except Exception:
            pass
        db.delete(att)

    # 2. Delete messages next
    messages = db.exec(
        select(ChatMessageTable).where(ChatMessageTable.session_id == row.id)
    ).all()
    for m in messages:
        db.delete(m)

    row.updated_at = datetime.utcnow()
    db.add(row)
    db.commit()


@router.delete(
    "/sessions/{session_id}/attachments/{attachment_id}",
    status_code=status.HTTP_204_NO_CONTENT,
)
async def delete_attachment(
    session_id: UUID,
    attachment_id: UUID,
    current_user: TokenData = Depends(get_current_user),
    db: Session = Depends(get_session),
):
    row = _get_session_normalized(db, session_id)
    if not row:
        raise HTTPException(status_code=404, detail="Session not found")
    _assert_session_owner(row, current_user)
    att = db.get(ChatAttachmentTable, str(attachment_id))
    if not att or att.session_id != row.id:
        raise HTTPException(status_code=404, detail="Attachment not found")
    try:
        os.remove(att.file_path)
    except Exception:
        pass
    db.delete(att)
    db.commit()


class ProviderPingRequest(BaseModel):
    provider: str
    api_key: Optional[str] = None
    base_url: Optional[str] = None
    model: Optional[str] = None


@router.post("/providers/ping")
async def ping_provider(req: ProviderPingRequest):
    import httpx

    provider = req.provider.lower()
    api_key = req.api_key or ""
    base_url = req.base_url or ""
    model = req.model or ""

    try:
        if provider == "lmstudio":
            candidate_bases = build_local_provider_base_candidates(base_url)
            model_name = model or "liquid/lfm2.5-1.2b"
            headers = {"Content-Type": "application/json"}
            payload = {
                "model": model_name,
                "max_tokens": 1,
                "messages": [{"role": "user", "content": "ping"}],
            }
        elif provider == "opencode":
            endpoint = (
                f"{base_url.rstrip('/')}/chat/completions"
                if base_url
                else "https://opencode.ai/zen/v1/chat/completions"
            )
            model_name = model or "gpt-5.5"
            headers = {"Content-Type": "application/json"}
            if api_key:
                headers["Authorization"] = f"Bearer {api_key}"
            payload = {
                "model": model_name,
                "max_tokens": 1,
                "messages": [{"role": "user", "content": "ping"}],
            }
        elif provider == "openai":
            endpoint = "https://api.openai.com/v1/chat/completions"
            model_name = model or "gpt-4o-mini"
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            }
            payload = {
                "model": model_name,
                "max_tokens": 1,
                "messages": [{"role": "user", "content": "ping"}],
            }
        elif provider == "groq":
            endpoint = "https://api.groq.com/openai/v1/chat/completions"
            model_name = model or "llama-3.1-70b-versatile"
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            }
            payload = {
                "model": model_name,
                "max_tokens": 1,
                "messages": [{"role": "user", "content": "ping"}],
            }
        elif provider == "claude":
            endpoint = "https://api.anthropic.com/v1/messages"
            model_name = model or "claude-3-5-sonnet-20241022"
            headers = {
                "Content-Type": "application/json",
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            }
            payload = {
                "model": model_name,
                "max_tokens": 1,
                "messages": [{"role": "user", "content": "ping"}],
            }
        elif provider == "gemini":
            endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{model or 'gemini-1.5-pro'}:generateContent?key={api_key}"
            headers = {"Content-Type": "application/json"}
            payload = {
                "contents": [{"parts": [{"text": "ping"}]}],
                "generationConfig": {"maxOutputTokens": 1},
            }
        else:
            return {
                "status": "unsupported",
                "message": f"Unsupported provider: {provider}",
            }

        async with httpx.AsyncClient(timeout=8.0) as client:
            if provider == "lmstudio":
                last_error = None
                for endpoint in [
                    f"{candidate.rstrip('/')}/chat/completions"
                    for candidate in candidate_bases
                ]:
                    try:
                        resp = await client.post(
                            endpoint, json=payload, headers=headers
                        )
                        if resp.status_code == 200:
                            return {
                                "status": "healthy",
                                "message": "Connection healthy! Model is fully responsive.",
                            }
                        last_error = f"Server returned error code {resp.status_code}: {resp.text[:200]}"
                    except Exception as exc:
                        last_error = str(exc)
                return {
                    "status": "offline",
                    "message": f"Ping connection failed: {last_error or 'All connection attempts failed'}",
                }

            resp = await client.post(endpoint, json=payload, headers=headers)
            if resp.status_code == 200:
                return {
                    "status": "healthy",
                    "message": "Connection healthy! Model is fully responsive.",
                }
            return {
                "status": "error",
                "message": f"Server returned error code {resp.status_code}: {resp.text[:200]}",
            }
    except Exception as e:
        return {"status": "offline", "message": f"Ping connection failed: {str(e)}"}


class ProviderDiscoverRequest(BaseModel):
    provider: str
    api_key: Optional[str] = None
    base_url: Optional[str] = None


@router.post("/providers/discover")
async def discover_provider_models(req: ProviderDiscoverRequest):
    import httpx

    provider = req.provider.lower()
    api_key = req.api_key or ""
    base_url = req.base_url or ""

    try:
        if provider in [
            "lmstudio",
            "openai",
            "groq",
            "custom",
            "opencode",
            "openai-compatible",
        ]:
            if provider == "lmstudio":
                candidate_bases = build_local_provider_base_candidates(base_url)
            elif provider == "openai":
                url = "https://api.openai.com/v1/models"
            elif provider == "groq":
                url = "https://api.groq.com/openai/v1/models"
            elif provider == "opencode":
                url = (
                    f"{base_url.rstrip('/')}/models"
                    if base_url
                    else "https://opencode.ai/zen/v1/models"
                )
            else:
                if not base_url.strip():
                    return {
                        "status": "error",
                        "message": "Base URL is required for OpenAI-compatible providers",
                        "models": [],
                    }
                url = f"{base_url.rstrip('/')}/models"

            headers = {"Content-Type": "application/json"}
            if api_key:
                headers["Authorization"] = f"Bearer {api_key}"

            async with httpx.AsyncClient(timeout=6.0) as client:
                if provider == "lmstudio":
                    last_error = None
                    for url in [
                        f"{candidate.rstrip('/')}/models"
                        for candidate in candidate_bases
                    ]:
                        try:
                            resp = await client.get(url, headers=headers)
                            if resp.status_code == 200:
                                data = resp.json()
                                model_list = [
                                    m["id"] for m in data.get("data", []) if "id" in m
                                ]
                                return {"status": "success", "models": model_list}
                            last_error = f"Server status {resp.status_code}"
                        except Exception as exc:
                            last_error = str(exc)
                    return {
                        "status": "error",
                        "message": last_error or "Live discovery failed",
                        "models": [],
                    }

                resp = await client.get(url, headers=headers)
                if resp.status_code == 200:
                    data = resp.json()
                    model_list = [m["id"] for m in data.get("data", []) if "id" in m]
                    return {"status": "success", "models": model_list}
                return {
                    "status": "error",
                    "message": f"Server status {resp.status_code}",
                    "models": [],
                }

        elif provider == "ollama":
            url = (
                f"{base_url.rstrip('/')}/api/tags"
                if base_url
                else "http://127.0.0.1:11434/api/tags"
            )
            async with httpx.AsyncClient(timeout=6.0) as client:
                resp = await client.get(url)
                if resp.status_code == 200:
                    data = resp.json()
                    model_list = [
                        m["name"] for m in data.get("models", []) if "name" in m
                    ]
                    return {"status": "success", "models": model_list}
                else:
                    return {
                        "status": "error",
                        "message": f"Server status {resp.status_code}",
                        "models": [],
                    }

        elif provider in ["google", "gemini"]:
            url = (
                f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
            )
            async with httpx.AsyncClient(timeout=6.0) as client:
                resp = await client.get(url)
                if resp.status_code == 200:
                    data = resp.json()
                    model_list = [
                        m["name"].split("/")[-1]
                        for m in data.get("models", [])
                        if "name" in m
                    ]
                    return {"status": "success", "models": model_list}
                else:
                    return {
                        "status": "error",
                        "message": f"Server status {resp.status_code}",
                        "models": [],
                    }

        elif provider in ["anthropic", "claude"]:
            return {
                "status": "success",
                "models": ["claude-3-5-sonnet-20241022", "claude-3-haiku-20240307"],
            }

        else:
            return {
                "status": "unsupported",
                "message": "Discovery not supported",
                "models": [],
            }

    except Exception as e:
        return {"status": "exception", "message": str(e), "models": []}
